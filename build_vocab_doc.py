#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成期末备考词汇与提纲 Word 文档"""

import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ===== 样式设置 =====
style = doc.styles['Normal']
font = style.font
font.name = '等线'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')

# ===== 标题 =====
title = doc.add_heading('2026春 读写II 期末备考资料', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    '涵盖范围：《综4》Unit 1, 4, 5, 6 (Text A) + 《原生态》蓝皮书 Unit 7, 8, 10, 12 (Passage I)'
).alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('=' * 60)

# ==========================================
# 第一部分：备考提纲
# ==========================================
doc.add_heading('第一部分：备考提纲', level=1)

doc.add_heading('一、题型与分值（考试时间90分钟）', level=2)

# 题型表格
table = doc.add_table(rows=7, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['大题', '题型', '题量/分值', '备考策略']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True

data = [
    ['第一大题\n(35分)',
     '1.1 课内词汇释义\n1.2 课外词汇释义',
     '30题×1分=30分\n10题×0.5分=5分',
     '★重点：标蓝词汇！30题全部来自标蓝词汇\n课外题难度六级，每题15秒'],
    ['第二大题\n(40分)',
     '2.1 传统篇章阅读(3篇)\n2.2 七选五标题匹配(1篇)',
     '—\n—',
     'Passage 3难度最大；标题匹配不亚于Task 2学术文章\n第一段和最后一段不选标题'],
    ['第三大题\n(25分)',
     '3.1 课内完形填空\n3.2 课外完形填空',
     '10题×1分=10分\n15选10=15分',
     '课内：来自Unit1/4/5/6 Text A 任一篇(3分钟搞定)\n课外：CET6形式'],
]

for r, row_data in enumerate(data):
    for c, text in enumerate(row_data):
        table.rows[r+1].cells[c].text = text

doc.add_paragraph()

doc.add_heading('二、复习优先级', level=2)

p = doc.add_paragraph()
p.add_run('第一优先：').bold = True
p.add_run('《综4》Unit 1, 4, 5, 6 三篇 Text A 的标蓝词汇（课文与词汇讲解文档中标有 Text A 的部分）')

p = doc.add_paragraph()
p.add_run('第二优先：').bold = True
p.add_run('《原生态》蓝皮书 Unit 7, 8, 10, 12 四篇 Passage I 的标蓝词汇')

p = doc.add_paragraph()
p.add_run('第三优先（如有时间）：').bold = True
p.add_run('标有 Text B 的课文与词汇讲解')

doc.add_heading('三、时间分配建议', level=2)

time_table = doc.add_table(rows=5, cols=2)
time_table.style = 'Light Grid Accent 1'
time_table.rows[0].cells[0].text = '模块'
time_table.rows[0].cells[1].text = '建议时间'
for p in time_table.rows[0].cells[0].paragraphs:
    for run in p.runs:
        run.bold = True
for p in time_table.rows[0].cells[1].paragraphs:
    for run in p.runs:
        run.bold = True

time_data = [
    ['四篇阅读', '60分钟内（每篇8-15分钟）'],
    ['课外完形填空', '17分钟'],
    ['课内完形填空', '3分钟'],
    ['40道选择题', '10分钟（每题15秒）'],
]
for r, row_data in enumerate(time_data):
    time_table.rows[r+1].cells[0].text = row_data[0]
    time_table.rows[r+1].cells[1].text = row_data[1]

doc.add_paragraph()

doc.add_heading('四、考试覆盖文章', level=2)

articles = [
    ('《综4》Unit 1 Text A', 'The Icy Defender (Nila B. Smith)', '拿破仑与希特勒侵俄，俄国寒冬'),
    ('《综4》Unit 4 Text A', 'In Search of Davos Man (Peter Gumbel)', '全球化、达沃斯人、国家认同'),
    ('《综4》Unit 5 Text A', 'A Friend in Need (Somerset Maugham)', '伯顿的故事、人性黑暗面'),
    ('《综4》Unit 6 Text A', 'Old Father Time Becomes a Terror (Richard Tomkins)', '科技与时间压力'),
    ('蓝皮书 Unit 7 Passage I', 'Why We Should Study Cancer Like We Study Ecosystems', '癌症生态学视角'),
    ('蓝皮书 Unit 8 Passage I', 'Water Damage (Martha Southgate)', '非裔美国人游泳问题'),
    ('蓝皮书 Unit 10 Passage I', 'The Trouble With Online Education (Mark Edmundson)', '在线教育之弊'),
    ('蓝皮书 Unit 12 Passage I', 'Why Waiting Is Torture (Alex Stone)', '排队心理学'),
]

art_table = doc.add_table(rows=len(articles)+1, cols=3)
art_table.style = 'Light Grid Accent 1'
art_table.rows[0].cells[0].text = '来源'
art_table.rows[0].cells[1].text = '标题/作者'
art_table.rows[0].cells[2].text = '主题'
for p in art_table.rows[0].cells[0].paragraphs:
    for run in p.runs:
        run.bold = True
for p in art_table.rows[0].cells[1].paragraphs:
    for run in p.runs:
        run.bold = True
for p in art_table.rows[0].cells[2].paragraphs:
    for run in p.runs:
        run.bold = True

for r, (src, title, theme) in enumerate(articles):
    art_table.rows[r+1].cells[0].text = src
    art_table.rows[r+1].cells[1].text = title
    art_table.rows[r+1].cells[2].text = theme

doc.add_page_break()

# ==========================================
# 第二部分：标蓝词汇（按单元整理，方便背诵）
# ==========================================
doc.add_heading('第二部分：标蓝词汇背诵手册', level=1)
doc.add_paragraph('说明：所有词汇均来自老师上传的"课文与词汇讲解"文档中的标蓝词汇（罗马数字标注）。30道课内词汇题全部出自以下词汇，无一例外。')
doc.add_paragraph('背诵提示：★ 标记为重点词汇，常考近义词/词形辨析/短语搭配。')

# ===== 词汇数据 =====

vocab_units = []

# ---- Unit 1 ----
unit1_vocab = [
    ('fierce', 'adj.', '顽强的，猛烈的', 'intense, powerful, violent', 'fierce resistance 顽强抵抗'),
    ('devastating', 'adj.', '极具摧毁力的', 'highly destructive, extremely powerful', 'a devastating enemy'),
    ('raw', 'adj.', '阴冷潮湿的（罕见用法）', 'unpleasantly cold or damp', 'raw winter; 本义"生肉"'),
    ('bitter', 'adj.', '严酷的，刺骨的', 'raw, harsh, cold', 'bitter winter'),
    ('bleak', 'adj.', '阴冷的，荒凉的', 'raw, bleak, cold', 'bleak winter ★三词近义连用'),
    ('launch', 'v.', '发动，发起，投放', 'start, release, initiate', 'launch a war/campaign/attack/project/product'),
    ('might', 'n.', '力量，威力', 'power, strength', 'military might; adj.=mighty; adv.=mightily'),
    ('unequaled', 'adj.', '无与伦比的', 'unparalleled, incomparable, unrivaled', 'unequaled military might'),
    ('mow down', 'phr.', '残忍地大量消灭', 'kill or destroy in great numbers', 'mowed down resistance'),
    ('conquest', 'n.', '征服；被征服之地', 'act of conquering', 'v.=conquer'),
    ('bog down', 'phr.', '陷入困境，停滞', 'cause to sink; get stuck', 'advance became bogged down'),
    ('engage', 'v.', '吸引；交火；订婚', '1)attract 2)enter combat 3)pledge to marry', '★三含义必考其一'),
    ('take the gamble', 'phr.', '冒险一试', 'do something risky', 'took the gamble of pressing on'),
    ('press on to', 'phr.', '大军压境', 'advance toward', 'pressed on to Moscow'),
    ('raging', 'adj.', '猛烈的，烈火般的', 'violent, forceful, severe', 'a raging fire 烈火'),
    ('bide one\'s time', 'phr.', '等待时机', 'wait for further deployment', '固定搭配'),
    ('trap', 'v.', '困住，阻止逃脱', 'prevent from escaping', 'nearly trapped the retreating French'),
    ('a stroke of luck', 'phr.', '一丝运气', 'a bit of fortune', 'a stroke of humor/madness'),
    ('alliance', 'n.', '联合（抽象）', 'state of being allied', 'ally n./v. 盟友/联合（具象）'),
    ('abdicate', 'v.', '放弃王位，退位', 'renounce one\'s throne', 'Napoleon abdicated'),
    ('exile', 'n.', '流放', 'forced absence from home', 'went into exile'),
    ('campaign', 'n.', '战役；活动', 'series of military operations', '★区分: campaign/champaign/champagne'),
    ('catch sb off guard', 'phr.', '猝不及防', 'surprise sb unexpectedly', 'was caught off guard by the invasion'),
    ('render', 'v.', '致使；翻译', '1)make 2)translate', '★两含义: render useless; render into English'),
    ('desperate', 'adj.', '深感绝望的', 'having lost hope', 'the situation became desperate'),
    ('settle over', 'phr.', '逐渐降临', 'sink or descend gradually', 'winter settled over the Soviet Union'),
    ('liberate', 'v.', '解放', 'set free', 'PLA = People\'s Liberation Army'),
    ('unspeakable', 'adj.', '无法言说的（糟）', 'beyond description, inexpressibly bad', '★区分: unspeakable vs. speechless'),
    ('the elements', 'n.(pl.)', '恶劣天气（特殊用法）', 'violent or severe weather', '★必须加s，不表示"元素"'),
    ('take a toll on', 'phr.', '带来损害', 'cause cumulative negative effect', 'Snow took its toll on armies.'),
]

vocab_units.append(('《综4》Unit 1 Text A — The Icy Defender', unit1_vocab))

# ---- Unit 4 ----
unit4_vocab = [
    ('assets', 'n.(pl.)', '资产', 'property (pl.); advantage (sing.)', '★复数=资产; 单数=优势: His wit is his chief asset.'),
    ('fervent', 'adj.', '热情的，强烈的', 'passionate, zealous, ardent, keen', 'a fervent believer'),
    ('keen', 'adj.', '敏锐的；热切的；极好的', 'sharp; eager; excellent', '★高频词: keen to see; a keen production'),
    ('reckon', 'v.', '估计，估摸', 'estimate, count', 'reckon the height; ★也表consider'),
    ('globetrotting', 'n./adj.', '环球旅行', 'globe+trot; travel worldwide', '碎步慢跑=trot≈jog'),
    ('identify', 'v.', '确立身份；等同于', 'recognize identity; associate with', '★identify with 感同身受; identical 相同的'),
    ('implications', 'n.(pl.)', '潜在可能性/意义', 'possibilities, significance', 'v.=imply 暗示'),
    ('commit', 'v.', '做(错事)；奉献', '1)perform(wrong) 2)bind/devote', '★commit a crime; be committed to nation'),
    ('endorse', 'v.', '公开支持/认可', 'support or approve publicly', 'endorse a global outlook'),
    ('outlook', 'n.', '观点，看法', 'view, viewpoint', 'global outlook'),
    ('content', 'adj./n.', '满足的/内容', 'satisfied (重音在后); substance (重音在前)', '★区分: consent赞同 / content满足 / contend争辩 / contest竞争'),
    ('beat a path to one\'s door', 'phr.', '急切寻求合作', 'be eager to do business with', 'beat a path to China\'s door'),
    ('overseas', 'adj./adv.', '海外', 'across the sea', 'oversea(英)/overseas(美) 皆可'),
    ('graduate', 'v.', '授予学位', 'grant an academic diploma', '★及物用法: China graduates engineers.'),
    ('projections', 'n.(pl.)', '预判，估计', 'estimate of future based on trend', '≈predictions, estimates'),
    ('aggregate', 'n./v.', '集合/聚集', 'collection; gather', '反义词: segregate 分离'),
    ('tension', 'n.', '紧张', 'nervous strain', '★派生: tensed (书面语)=nervous'),
    ('aspire', 'v.', '努力追求', 'strive for (不及物+to)', 'aspire to a career; ★区分: acquire获得 / inspire激励'),
]
vocab_units.append(('《综4》Unit 4 Text A — In Search of Davos Man', unit4_vocab))

# ---- Unit 5 ----
unit5_vocab = [
    ('vanity', 'n.', '虚荣；自负', '1)emptiness 2)inflated pride', '本文取"自负"义'),
    ('all of a piece', 'phr.', '表里如一的人', 'a consistent person', 'Here if ever was a man all of a piece.'),
    ('in accordance with', 'phr.', '与……一致', 'consistent with', '反义: be at odds with 与……不一致'),
    ('sensible', 'adj.', '理智的；合理的', 'reasonable, rational', '★sense派生: sensual/sensory(感官); sensitive(敏感); sensational(轰动的)'),
    ('tell with point', 'phr.', '重点明确地讲述', 'tell a story effectively', 'good and spicy story 饶有趣味的故事'),
    ('arouse', 'v.', '激起', 'stir, raise', 'arouse one\'s instincts of protection'),
    ('could not bear to', 'phr.', '不忍心做', 'cannot endure to', 'could not bear to hurt a fly'),
    ('uncanny', 'adj.', '不可思议的，神奇的', 'beyond normal, unusual', 'It was uncanny. 这简直神了。'),
    ('chuckle', 'n./v.', '咯咯笑（捂嘴）', 'laugh quietly', '★区分: gurgling咕噜声; tickling/giggling吱吱笑; belly laugh捧腹大笑; horselaugh傻笑'),
    ('get trimmed', 'phr.', '被骗', 'get cheated (trim的罕见用法)', 'He\'d got trimmed at poker.'),
    ('drive at', 'phr.', '想要表达', 'hope to express', 'I got some glimmering of what he was driving at.'),
    ('little tin gods', 'phr.', '自命不凡的人', 'minor people feeling self-important', 'tin锡=廉价金属'),
    ('funk', 'v.', '因畏惧而放弃', 'give up out of fear', 'Did he funk it?'),
    ('dissipation', 'n.', '放荡，放任', 'wildness, indulgence', 'ruined his constitution by drink and dissipation'),
    ('candid', 'adj.', '真诚的，坦率的', 'honest, sincere', 'kind and candid blue eyes'),
]
vocab_units.append(('《综4》Unit 5 Text A — A Friend in Need', unit5_vocab))

# ---- Unit 6 ----
unit6_vocab = [
    ('waste away', 'phr.', '变虚弱，消磨', 'become weaker', 'waste away in idleness'),
    ('idleness', 'n.', '无所事事', 'laziness', 'idle=lazy'),
    ('enslave', 'v.', '奴役', 'en+slave: cause to be a slave', '★en-前缀: endear/enslave/enable'),
    ('as...as', '结构', '同等比较', 'equally...like...', 'as many arrive in a year as once arrived in a millennium'),
    ('be/get stuck in', 'phr.', '被困住', 'be trapped', '★非正式，不用于学术写作'),
    ('undertake', 'v.', '承担；着手', '1)take obligation 2)begin to do', '★高级词; undertaking=task'),
    ('trans-', '前缀', '跨越', 'across', 'transatlantic, transcultural, transcontinental'),
    ('fraction', 'n.', '部分', 'portion, part', '★辨析: faction党派/fiction小说/fracture骨折/friction摩擦'),
    ('proliferate', 'v.', '激增，扩散', 'multiply, increase rapidly', 'Magazines proliferate.'),
    ('scholarship', 'n.', '奖学金；学术/知识库', '1)grant 2)fund of knowledge', '★两含义均常考'),
    ('amount to', 'phr.', '等同于', 'be equal to, be the same as', 'what amounts to a minute proportion'),
    ('oblige', 'v.', '效劳；迫使', 'do as a favour; compel', 'we do our best to oblige'),
    ('abundant/abundance', 'adj./n.', '富足，充裕', 'full of variety, ample', '★辨析: abandon 摒弃/abandoned 被遗弃的'),
    ('shorthand', 'n.', '速记；简化的说法', 'method of quick writing; simplified expression', 'a convenient shorthand'),
    ('peripherally', 'adv.', '边缘化地', 'marginally, slightly', 'peripherally involved'),
    ('impose', 'v.', '强加', 'force sth unwelcome to be accepted', '★六级词汇: self-imposed pressure'),
    ('nurture', 'v./n.', '教育，培养', 'educate, develop', 'nature vs. nurture 先天vs后天'),
    ('appalling', 'adj.', '令人震惊的', 'shocking, startling, astonishing', '=staggering(全学过)'),
    ('provoke', 'v.', '激起，煽动', 'stir, trigger', 'provoke reactions'),
    ('instant gratification', 'phr.', '即时满足', 'forego future benefit for immediate reward', '反义: delayed gratification 延迟满足'),
    ('streamline', 'v.', '使简化高效', 'make simpler or more efficient', 'streamline our lives'),
    ('arise', 'v.', '来源于', 'originate from a source', 'stress arises from surfeit'),
    ('cram', 'v.', '塞，灌，挤', 'fill, stuff', 'cram things into time'),
    ('an array of', 'phr.', '大量的，一系列', 'a variety of, a great number of', 'an array of choices'),
    ('confine', 'n./v.', '围栏；限制', 'border, wall; limit, restrict', 'within the confines of their village'),
    ('be doomed to', 'phr.', '注定要', 'be destined to, make certain the failure of', 'be doomed to mounting despair'),
    ('mounting', 'adj.', '不断累积的', 'increasing, growing, accumulating', 'mounting despair'),
]
vocab_units.append(('《综4》Unit 6 Text A — Old Father Time Becomes a Terror', unit6_vocab))

# ---- Unit 7 (蓝皮书) ----
unit7_vocab = [
    ('throw...out of balance', 'phr.', '打破平衡', 'make sth unsteady', 'threw a delicate ecosystem out of balance'),
    ('damaging', 'adj.', '有破坏力的', 'destructive', 'damaging invasive species'),
    ('angle', 'n.', '角度', 'aspect, side, facet', '★辨析: angle角度/angel天使/anger愤怒'),
    ('notorious', 'adj.', '臭名昭著的', 'widely known for bad quality', 'one of the most notorious maladies'),
    ('intensely', 'adv.', '深入地', 'profoundly, deeply', 'studied intensely'),
    ('localise', 'v.', '定位', 'situate, locate, position', 'localised to specific cells'),
    ('span', 'v.', '跨越', 'spread, extend wide across', 'spans the entire globe'),
    ('disruption', 'n.', '扰乱', 'interruption, disturbance, chaos', 'a disruption in the balance'),
    ('eat one\'s way through', 'phr.', '一路吃过去', 'consume progressively', '★类比: dance one\'s way to; study one\'s way to; run one\'s way to'),
    ('outbreak', 'n.', '爆发', 'sudden occurrence', 'the outbreak of a novel disease'),
    ('throw a monkey wrench into', 'phr.', '打破平静，使混乱', 'disrupt, cause problems to', 'cancer throws a monkey wrench into...'),
    ('otherwise', 'adv.', '在其他情况下', 'under other circumstances', 'an otherwise placid system'),
    ('insight', 'n.', '洞见，新理解', 'new way of understanding', 'provide insights'),
    ('ultimately', 'adv.', '最终', 'eventually, finally, in the end, at last', 'make cancer more treatable'),
    ('turn out', 'phr.', '结果是', 'be found to be', 'It turns out that...'),
    ('in terms of', 'phr.', '从……角度看', 'regarding, concerning, on the basis of', 'think of cancer in terms of mutated cells'),
    ('voracious', 'adj.', '贪婪的', 'greedy, having huge appetite', 'a voracious cancer cell'),
    ('extinct', 'adj.', '灭绝的', 'no longer active or existing', '★辨析: extinguish灭火/distinguish区别'),
    ('capacity', 'n.', '能力', 'ability, capability', 'the capacity to adapt'),
    ('painstakingly', 'adv.', '费力地', 'with great effort', '★拼写: 不是paintaking'),
    ('rule out', 'phr.', '排除', 'remove from consideration, exclude', 'rules out normal experiments'),
    ('go awry with', 'phr.', '走偏，出错', 'go away from the correct course', 'experiments could go awry'),
    ('suggest', 'v.', '暗示；建议', '1)indicate/imply 2)propose', '★两含义: 暗示(跟陈述语气); 建议(跟虚拟语气)'),
    ('denizen', 'n.', '居民，栖息者', 'resident, inhabitant', 'ecosystem denizens'),
    ('optimise', 'v.', '最优化', 'make as perfect/effective as possible', 'optimising the bottom line'),
    ('the bottom line', 'phr.', '最重要的事', 'the most important point', '★"坚守底线" ≠ stick to the bottom line'),
    ('tasty', 'adj.', '美味的', 'having a pleasing flavor', '≈juicy 美味多汁'),
    ('wild beast', 'n.', '野兽', 'wild animal', '兽性: bestiality'),
    ('victorious', 'adj.', '胜利的', 'successful', 'adj. of victory'),
    ('meaty', 'adj.', '肥硕的', 'heavily fleshed', 'meaty spoils'),
    ('indiscriminately', 'adv.', '不分青红皂白地', 'without discrimination', 'indiscriminately killing tumor cells'),
    ('longtime/long-term', 'adj.', '长久/影响持久的', 'long duration vs. lasting effect', '★longtime friend 老朋友; long-term effect 长远影响'),
    ('exclusively', 'adv.', '专门地，排他地', 'solely, simply, only', 'based exclusively on killing'),
    ('temporary', 'adj.', '短暂的', 'lasting for a very short time', '★辨析: contemporary现代的/当代的'),
    ('exceed', 'v.', '超过', 'surpass, outperform, transcend', 'the original number will be exceeded'),
    ('intriguing', 'adj.', '引人入胜的', 'engaging, attention-absorbing', 'intriguing implications'),
    ('irreparably', 'adv.', '不可逆转地', 'irreversibly', 'repar(修补)→reparable→irreparably'),
    ('proactive', 'adj.', '有远见的，前瞻的', 'acting in anticipation of future needs', '★pro-前缀=forward: foreseeing, farsighted'),
    ('manageable', 'adj.', '可处理的', 'capable of being handled/treated', 'until the cancer is manageable'),
    ('invalidate', 'v.', '使无效', 'make invalid/incorrect', 'does not invalidate but complements'),
    ('complement', 'v.', '补充，使完整', 'make complete', '★辨析: compliment褒奖/comprehensive全面/comprehend理解'),
    ('cannot', 'aux.', '不能', 'can not (必须连写!)', '★DON\'T write "can not"; 正确: cannot'),
]
vocab_units.append(('蓝皮书 Unit 7 Passage I — Why We Should Study Cancer Like We Study Ecosystems', unit7_vocab))

# ---- Unit 8 (蓝皮书) ----
unit8_vocab = [
    ('obscure', 'v.', '隐藏，掩盖', 'hide, cover, conceal', 'performances obscure a disturbing truth'),
    ('anomaly', 'n.', '反常，罕见', 'something different, abnormal', 'anomalies in elite levels'),
    ('rudimentary', 'adj.', '基本的', 'basic, fundamental, elementary', 'rudimentary swimming skills'),
    ('evangelist', 'n.', '鼓吹者，宣扬者', 'enthusiastic advocate (小写)', '★大写E=福音书作者; 小写=事业宣扬者'),
    ('initiative', 'n.', '计划，倡议', 'a plan, a cause', 'a water-safety initiative'),
    ('give away', 'phr.', '赠予', 'give as a gift for', 'give swimming lessons away'),
    ('mandatory', 'adj.', '强制的', 'required by law or rule, obligatory', 'made swimming lessons mandatory'),
    ('talk sb into', 'phr.', '说服某人做', 'persuade to do', '★into有被动语义: Don\'t talk yourself into staying.'),
    ('bear out', 'phr.', '证实', 'confirm', 'This bears out the finding.'),
    ('respondent', 'adj.', '做出回应的', 'giving a response', 'across all respondent race groups'),
    ('take an ugly toll', 'phr.', '带来惨痛代价', 'take an unpleasant cost', 'segregation took its ugly toll'),
    ('motel', 'n.', '汽车旅馆', 'motor + hotel', '位于城郊的简易旅馆'),
    ('stage', 'v.', '上演', 'put on, perform', 'stage a swim-in'),
    ('swim-in/sit-in/study-in', 'n.', '游泳/静坐/学习抗议', 'protest by swimming/sitting/studying', '★-in后缀=protest'),
    ('institutionalised', 'adj.', '根深蒂固的', 'established as common belief', 'institutionalised racism'),
    ('shore up', 'phr.', '用证据支持', 'support with evidence', 'was shored up by specious scholarship'),
    ('instill', 'v.', '灌输', 'implant gradually', 'fear instilled in African-Americans; ★≈stuff/squeeze knowledge into'),
    ('self-perpetuating', 'adj.', '自我永续的', 'causing oneself to continually follow', 'the fear has become self-perpetuating'),
    ('be blessed with', 'phr.', '有幸拥有', 'be endowed with, be favoured with', 'This country is blessed with...'),
    ('regardless of', 'phr.', '撇开……不论', 'without considering', 'regardless of its racial makeup'),
    ('makeup', 'n.', '构成', 'composition', 'racial makeup 种族构成'),
    ('shuttle', 'v./n.', '穿梭运输/班车', 'transport back and forth', 'shuttling students to swim programs'),
    ('partner A with B', 'phr.', '将A与B配对', '名词动用: pair A with B', '★类比: room with sb; spoon me soup; eye her; bicycle home'),
]
vocab_units.append(('蓝皮书 Unit 8 Passage I — Water Damage', unit8_vocab))

# ---- Unit 10 (蓝皮书) ----
unit10_vocab = [
    ('cringe', 'v.', '尴尬，难为情', 'feel embarrassed', 'This line always makes me cringe.'),
    ('latent', 'adj.', '潜在的，不明显的', 'present but not obvious', 'latent in the sentiment is an important truth'),
    ('sentiment', 'n.', '观点，意见', 'opinion, idea', '此处=opinion'),
    ('come close to', 'phr.', '差一点就', 'almost do sth', 'came close to tearing the university apart'),
    ('summarily', 'adv.', '迅速地，即刻地', 'quickly without delay', 'summarily dismissed'),
    ('ever', 'adv.', '（强调）', 'emphasis: really', 'can online education ever be...?'),
    ('illuminating', 'adj.', '启发人的', 'formative, inspiring', 'the notion is illuminating'),
    ('intellectually', 'adv.', '在智力/思考层面', 'in terms of rational thought', 'where they are intellectually'),
    ('be worth doing', 'phr.', '值得做', 'deserve to do', 'Is it worth adding a film version?'),
    ('adept', 'adj.', '熟练的', 'proficient, expert, skillful', '★辨析: adapt适应/adopt实施/adept熟练'),
    ('rendering', 'n.', '呈现，塑造', 'representation', 'Shakespeare\'s rendering of character'),
    ('disciplining', 'adj.', '基于规则的', 'controlled, ordered', 'against that disciplining background'),
    ('superb', 'adj.', '极好的', 'having high degree of excellent', 'They are superb at sensing the mood.'),
    ('engaged', 'adj.', '专注的', 'concentrated, focused', 'when the class is engaged'),
    ('discern', 'v.', '察觉，识别', 'detect, recognise', 'discerning who is out there'),
    ('one-size-fits-all', 'adj.', '均码的；一刀切的', 'suitable for all', 'online education is a one-size-fits-all endeavor'),
    ('endeavor', 'n.', '活动，努力', 'activity, effort', 'a one-size-fits-all endeavor'),
    ('articulate', 'adj.', '口齿伶俐的', 'expressing clearly and effectively', 'splendidly articulate'),
    ('on hand / in hand / at hand', 'phr.', '在场/在掌控中/手边', 'present / in possession / within reach', '★辨析: 三个短语不同含义'),
    ('immediate', 'adj.', '直接的，实时的', 'current, direct (非"立刻")', 'create an immediate community'),
    ('vital', 'adj.', '充满生机的', 'animated, invigorating', 'a vital community of learning'),
    ('promise to do', 'phr.', '预示着（常贬义）', 'be expected to (often negative)', 'Internet learning promises to make life more sterile.'),
]
vocab_units.append(('蓝皮书 Unit 10 Passage I — The Trouble With Online Education', unit10_vocab))

# ---- Unit 12 (蓝皮书) ----
unit12_vocab = [
    ('inordinate', 'adj.', '过多的，超出预期的', 'exceeding reasonable limits', 'an inordinate number of complaints'),
    ('lodge', 'v.', '提出（投诉）', 'register, file', 'lodge complaints'),
    ('persist', 'v.', '持续不断', 'continue to exist', 'the complaints persisted'),
    ('on-site', 'adj./adv.', '现场的', 'at the place where event happens', 'on-site analysis'),
    ('route', 'v.', '按既定路线传送', 'send through a chosen direction', 'routed bags to the outermost carousel'),
    ('outermost', 'adj.', '最外头的', 'the most outer', 'outermost carousel'),
    ('carousel', 'n.', '行李转盘', 'baggage claim conveyor', '本义: 旋转木马 → 语义延伸→行李转盘'),
    ('define', 'v.', '决定，支配', 'determine, dictate, govern', 'the experience is defined by the wait length'),
    ('note', 'v.', '表明（研究结果）', 'indicate (observation result)', 'notes the researcher'),
    ('foremost', 'adj.', '最前沿的', 'at the most forefront, in the lead', 'the world\'s foremost expert'),
    ('rationale', 'n.', '深层原因', 'underlying reason', 'the rationale behind the mirrors'),
    ('slyly', 'adv.', '偷偷摸摸地', 'in a cunning manner', 'slyly ogle other passengers'),
    ('ogle', 'v.', '盯着看', 'eye with greedy attention', 'slyly ogle other passengers'),
    ('drudgery', 'n.', '无聊，单调', 'boredom, dullness', 'the drudgery of unoccupied time'),
    ('account for', 'phr.', '解释，是……的原因', 'explain, be the reason for', 'accounts for the popularity of impulse-buy items'),
    ('impulse-buy items', 'phr.', '冲动购买品', 'items bought impulsively', '口香糖、小报等排队区商品'),
    ('magnify', 'v.', '放大，夸张', 'exaggerate; enlarge', 'uncertainty magnifies the stress'),
    ('tenor', 'n.', '整体过程/基调', 'general way of process', 'the tenor of the experience'),
    ('beat', 'v.', '胜过，击败', 'defeat, outdo', 'beating expectations buoys our mood'),
    ('buoy', 'v.', '提振，支撑', 'support, keep afloat (名词动用)', '本义: 浮力 → 提振情绪'),
    ('retrospect/retrospective', 'n./adj.', '回顾/回顾性的', 'looking back', '反义: prospect 展望'),
    ('audit', 'n.', '回顾，审查', 'review', 'our retrospective audit'),
    ('skew', 'v.', '偏向（怀疑态度）', 'look with suspicion', 'skew toward cynicism'),
    ('cynicism', 'n.', '不满，挑剔', 'disapproval, complaint', 'skew toward cynicism'),
    ('opt for', 'phr.', '选择', 'choose, select', '★固定搭配，务必接for'),
    ('wrap', 'v.', '包裹，蜿蜒折叠', 'wind or fold to cover', '★本学期多次出现: wrapping paper; casings wrapped up; wrap around buildings'),
    ('serpentine', 'adj.', '蛇形的', 'resembling a serpent', 'serpentine queues'),
    ('deviation', 'n.', '偏离', 'departure from established norm', 'any deviation is a mark of iniquity'),
    ('iniquity', 'n.', '不公，邪恶', 'injustice, wickedness', 'a mark of iniquity'),
    ('stab', 'v.', '刺伤', 'pierce or wound with pointed weapon', 'a man was stabbed at a post office'),
    ('intrusion', 'n.', '侵犯', 'act of wrongfully entering', 'unwelcome intrusions'),
    ('slips', 'n.', '偷偷溜进队伍', 'cutting in the queue', '★排队用语'),
    ('skips', 'n.', '从后往前插队', 'jumping the line', '★排队用语'),
    ('implicit', 'adj.', '潜在的，不成文的', 'implied, not consciously recognised', '反义: explicit 明确的'),
    ('norms', 'n.', '标准，规范', 'a set of standards', 'an implicit set of norms'),
    ('as opposed to', 'phr.', '与……相对', 'indicating difference from', 'single-queue as opposed to multi-queue'),
    ('wind up', 'phr.', '最终……', 'come to an end after action', 'you wind up kicking yourself'),
    ('cognitive', 'adj.', '认知的', 'related to conscious intellectual activity', 'cognitive asymmetry'),
    ('symmetry/asymmetry', 'n.', '对称/不对称', 'balanced proportions / the opposite', 'a curious cognitive asymmetry'),
    ('fixate', 'v.', '专注于', 'focus on', 'fixate on the line they\'re losing to'),
    ('dictate', 'v.', '支配，要求', 'impose, determine, require', 'fairness dictates that...'),
    ('commensurate', 'adj.', '相称的，匹配的', 'equal in measure', 'be commensurate with the value'),
    ('express line', 'n.', '快捷通道', 'supermarket fast lane', '≤10-15件商品可用'),
    ('sanction', 'v.', '赞同，认可', 'approve of, give consent to', 'socially sanctioned violation'),
    ('nagging', 'adj.', '挑剔的，烦人的', 'persistently annoying', 'that nagging sensation'),
    ('dwindling', 'adj.', '不断减少的', 'shrinking, becoming less', 'dwindling leisure time'),
    ('inevitable', 'adj.', '不可避免的', 'incapable of being avoided', 'inevitable delays'),
    ('bearable', 'adj.', '可忍受的', 'capable of being endured', 'a touch more bearable'),
]
vocab_units.append(('蓝皮书 Unit 12 Passage I — Why Waiting Is Torture', unit12_vocab))

# ===== 生成词汇表 =====
for unit_title, vocab_list in vocab_units:
    doc.add_heading(unit_title, level=2)

    # 创建表格
    tbl = doc.add_table(rows=len(vocab_list)+1, cols=5)
    tbl.style = 'Light Grid Accent 1'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    col_headers = ['词汇/短语', '词性', '中文释义', '英文释义/近义词', '例句/备注']
    for i, h in enumerate(col_headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    # 数据行
    for r, (word, pos, cn, en, note) in enumerate(vocab_list):
        row = tbl.rows[r+1]
        row.cells[0].text = word
        row.cells[1].text = pos
        row.cells[2].text = cn
        row.cells[3].text = en
        row.cells[4].text = note
        # 设置字体大小
        for c in range(5):
            for p in row.cells[c].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)

    # 设置列宽
    for row in tbl.rows:
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(1.2)
        row.cells[2].width = Cm(2.5)
        row.cells[3].width = Cm(4.5)
        row.cells[4].width = Cm(5.5)

    doc.add_paragraph()  # 间距

# ===== 第三部分：重点词形/词义辨析汇总 =====
doc.add_page_break()
doc.add_heading('第三部分：重点词形/词义辨析速查表', level=1)
doc.add_paragraph('以下为各单元反复出现的易混词辨析，考前务必过一遍。')

discriminations = [
    ('campaign / champaign / champagne', '战役/活动 | 平原 | 香槟酒', 'Unit 1'),
    ('unspeakable / speechless', '无法言说的(糟) | 因惊讶而无语的', 'Unit 1'),
    ('ally / alliance', '盟友(具象) | 联合(抽象)', 'Unit 1'),
    ('conquest / conquer', '征服(n.) | 征服(v.)', 'Unit 1'),
    ('content / consent / contend / contest', '满足 | 赞同 | 争辩 | 竞争', 'Unit 4 ★★★'),
    ('sensual / sensory / sensitive / sensible / sensational', '肉欲的 | 感官的 | 敏感的 | 理智的 | 轰动的', 'Unit 5 ★★★'),
    ('chuckle / gurgling / tickling(giggling) / belly laugh / horselaugh', '捂嘴笑 | 咕噜笑 | 吱吱笑 | 捧腹大笑 | 傻笑', 'Unit 5'),
    ('fraction / faction / fiction / fracture / friction', '部分 | 党派 | 小说 | 骨折 | 摩擦', 'Unit 6 ★★★'),
    ('abundant/abundance / abandon/abandoned', '富足 | 摒弃/被遗弃的', 'Unit 6'),
    ('nature / nurture', '先天遗传 | 后天养成', 'Unit 6'),
    ('instant gratification / delayed gratification', '即时满足 | 延迟满足', 'Unit 6'),
    ('angle / angel / anger', '角度 | 天使 | 愤怒', 'Unit 7'),
    ('extinct / extinguish / distinguish', '灭绝的 | 灭火 | 区别', 'Unit 7 ★★★'),
    ('temporary / contemporary', '短暂的 | 现代的/当代的', 'Unit 7'),
    ('invalidate / complement / compliment / comprehensive / comprehend', '使无效 | 补充 | 褒奖 | 全面 | 理解', 'Unit 7 ★★★'),
    ('longtime / long-term', '时间长久 | 影响持久', 'Unit 7'),
    ('adapt / adopt / adept', '适应 | 实施 | 熟练的', 'Unit 10 ★★★'),
    ('on hand / in hand / at hand', '在场 | 在掌控中 | 手边', 'Unit 10'),
    ('implicit / explicit', '潜在的/不成文的 | 明确的', 'Unit 12'),
    ('retrospect / prospect', '回顾 | 展望', 'Unit 12'),
    ('symmetry / asymmetry', '对称 | 不对称', 'Unit 12'),
]

disc_table = doc.add_table(rows=len(discriminations)+1, cols=3)
disc_table.style = 'Light Grid Accent 1'
disc_table.rows[0].cells[0].text = '词汇组'
disc_table.rows[0].cells[1].text = '含义辨析'
disc_table.rows[0].cells[2].text = '来源'
for p in disc_table.rows[0].cells[0].paragraphs:
    for run in p.runs:
        run.bold = True
for p in disc_table.rows[0].cells[1].paragraphs:
    for run in p.runs:
        run.bold = True
for p in disc_table.rows[0].cells[2].paragraphs:
    for run in p.runs:
        run.bold = True

for r, (words, meaning, source) in enumerate(discriminations):
    disc_table.rows[r+1].cells[0].text = words
    disc_table.rows[r+1].cells[1].text = meaning
    disc_table.rows[r+1].cells[2].text = source

doc.add_paragraph()

# ===== 第四部分：写作注意事项 =====
doc.add_heading('第四部分：学术写作注意事项', level=1)

warnings = [
    '1. cannot 必须连写，不能写成 "can not"（Unit 7 重点强调）',
    '2. 避免使用 you\'ve got... 等非正式表达，学术写作中应使用 there is/there are（Unit 6 重点强调）',
    '3. get stuck in 是非正式用语，不用于学术写作（Unit 6）',
    '4. had done (过去完成时) 必须有明确的过去时间起止点，无把握时少用（Unit 6）',
    '5. suggest 表"暗示"时跟陈述语气；表"建议"时跟虚拟语气 (should) do（Unit 7）',
]
for w in warnings:
    doc.add_paragraph(w)

# ===== 保存 =====
output_path = r'D:\辰辰\first CC\读写II_期末备考词汇与提纲.docx'
doc.save(output_path)
print(f'文档已保存至: {output_path}')
print('完成!')
