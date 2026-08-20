#!/usr/bin/env python3
"""生成《谪仙人》历史音乐剧完整策划文档（v2：含台本+样曲+双演员方案）"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── 页面设置 ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── 样式 ──
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)
rPr = style.element.get_or_add_rPr()
rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:eastAsia'), '宋体')
rPr.append(rFonts)

for level, (size, bold, color_hex) in enumerate([
    (24, True, '1a1a2e'), (16, True, '1a1a2e'), (13, True, '2d3a4a'), (12, True, '3d5066')
], 1):
    h_style = doc.styles[f'Heading {level}']
    h_style.font.name = '微软雅黑'
    h_style.font.size = Pt(size)
    h_style.font.bold = bold
    h_style.font.color.rgb = RGBColor(*tuple(int(color_hex[i:i+2],16) for i in range(0,6,2)))
    h_style.paragraph_format.space_before = Pt(18 if level <= 2 else 12)
    h_style.paragraph_format.space_after = Pt(8)

def add_para(text, style_name='Normal', bold=False, italic=False, alignment=None,
             font_size=None, font_name=None, color=None, indent=None):
    p = doc.add_paragraph(style=style_name)
    if alignment is not None:
        p.alignment = alignment
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size:
        run.font.size = Pt(font_size)
    if font_name:
        run.font.name = font_name
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), font_name)
        rPr.append(rFonts)
    if color:
        run.font.color.rgb = color
    return p

def add_quote(text, color='4472C4'):
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent = Cm(1)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '12')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), color)
    pBdr.append(left)
    pPr.append(pBdr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EBF0FA' if color == '4472C4' else 'FFF3E0')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = '楷体'
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '楷体')
    rPr.append(rFonts)
    return p

def add_table_with_data(headers, rows, col_widths=None):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9.5)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()
    return table

def add_stage_direction(text):
    """添加舞台提示（斜体灰色）"""
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(f'【{text}】')
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.name = '楷体'
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '楷体')
    rPr.append(rFonts)
    return p

def add_dialogue(character, text):
    """添加人物对白"""
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent = Cm(0.5)
    run_name = p.add_run(f'{character}：')
    run_name.bold = True
    run_name.font.size = Pt(10.5)
    run_text = p.add_run(text)
    run_text.font.size = Pt(10.5)
    return p

def add_lyric_line(text, indent=1):
    """添加歌词行"""
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = '楷体'
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '楷体')
    rPr.append(rFonts)
    return p

# ═══════════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════════
for _ in range(5):
    doc.add_paragraph()

add_para('谪 仙 人', 'Normal', bold=True, font_size=40, font_name='楷体',
         color=RGBColor(0x1a, 0x1a, 0x2e), alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('—— 李白历史摇滚音乐剧 · 完整策划方案 ——', 'Normal', font_size=14,
         font_name='楷体', color=RGBColor(0x66, 0x66, 0x66), alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

add_para('校园舞台版 · 低配不减质', 'Normal', font_size=12,
         font_name='微软雅黑', color=RGBColor(0x44, 0x72, 0xC4), alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('含：完整台本 +《将进酒》样曲 + 双演员方案', 'Normal', font_size=11,
         color=RGBColor(0x88, 0x88, 0x88), alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
add_para(f'策划日期：{datetime.date.today().strftime("%Y年%m月%d日")}', 'Normal',
         font_size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('架构参考：摇滚莫扎特（Mozart, l\'opéra rock）', 'Normal',
         font_size=10, color=RGBColor(0x99, 0x99, 0x99), alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 目录
# ═══════════════════════════════════════════════════
doc.add_heading('目录', level=1)

toc = [
    '第一部分：策划总纲',
    '  一、为什么李白 = 中国的莫扎特',
    '  二、人物设定与角色配置（含双演员方案）',
    '  三、剧情架构（两幕制·120分钟）',
    '  四、主旋律设计：《大鹏赋》',
    '  五、完整歌曲清单（18首）',
    '  六、音乐风格融合方案',
    '  七、女性角色的独立弧光',
    '',
    '第二部分：校园落地方案',
    '  八、校园舞台核心原则',
    '  九、舞美方案：一轮月亮+六扇屏风',
    '  十、乐队配置：五人成军',
    '  十一、场景落地方案（逐场详解）',
    '  十二、人员与预算',
    '  十三、排练时间线',
    '',
    '第三部分：完整舞台台本',
    '  十四、台本凡例',
    '  十五、第一幕台本（序曲→将进酒→幕间）',
    '  十六、第二幕台本（梁园吟→骑鲸→谢幕）',
    '',
    '第四部分：样曲',
    '  十七、《将进酒》完整歌词与编曲说明',
]
for item in toc:
    if item == '':
        doc.add_paragraph()
    elif item.startswith('第'):
        add_para(item, bold=True, font_size=11, font_name='微软雅黑')
    else:
        add_para(item, font_size=10.5)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 第一部分：策划总纲
# ═══════════════════════════════════════════════════
add_para('第一部分：策划总纲', 'Normal', bold=True, font_size=16, font_name='微软雅黑',
         color=RGBColor(0x1a, 0x1a, 0x2e), alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

# ── 一 ──
doc.add_heading('一、为什么李白 = 中国的莫扎特', level=1)

add_para('摇滚莫扎特成功的内核公式——天才的耀眼 + 体制的冲突 + 酒神精神 + 早逝的惋惜——在李白的生命中悉数存在，甚至更为极致。')

add_table_with_data(
    ['维度', '莫扎特', '李白'],
    [
        ['天赋', '音乐神童，即兴弹奏', '五岁诵六甲，十岁观百家，十五好剑术，二十文章惊海内'],
        ['与权力的关系', '为萨尔茨堡大主教服务，渴望自由', '供奉翰林，两年后主动求去——他选择了自由'],
        ['代表作遭遇', '《费加罗的婚礼》被审查', '因永王案以"附逆"罪下狱流放'],
        ['酒神精神', '热爱享乐、狂欢', '斗酒诗百篇——酒是他通往诗歌的钥匙'],
        ['结局', '35岁贫病而终，葬于无名公墓', '61岁病逝当涂——民间传说他投水追月'],
        ['父亲', '严格规划其音乐生涯', '李客，西域商人——给了他胡风、侠气、远方'],
        ['流浪', '巡回演出，欧洲各地', '一生仗剑去国，辞亲远游，足迹遍布大半个中国'],
        ['后世', '古典音乐之神', '诗仙——中国文学的最高峰'],
    ]
)

add_para('更为关键的是：莫扎特至死都在渴望自由而不得。李白选择了离开。他不是被放逐的天才——他是自我放逐的自由人。这使得李白比莫扎特更"摇滚"。', bold=True)

# ── 二 ──
doc.add_heading('二、人物设定与角色配置', level=1)

doc.add_heading('2.1 核心创新：双人分饰李白', level=2)

add_quote('🔥 全剧最重要的导演决策：李白由两位演员分饰——李白(A)演第一幕（20→42岁），李白(B)演第二幕（44→61岁）。这不是妥协，是升级。', 'E6550D')

add_para('李白在全剧中的年龄跨度为 40 年（20岁→61岁）。让同一个学生演员同时驾驭"佩剑出蜀的少年侠客"和"当涂病榻上垂死的六旬老人"，在校园选角条件下几乎不可能。而分为两人后：')

add_table_with_data(
    ['', '李白(A) —— 少年李白', '李白(B) —— 老年李白'],
    [
        ['饰演', '第一幕', '第二幕'],
        ['年龄', '20岁 → 42岁', '44岁 → 61岁'],
        ['声部', '摇滚男高音，清亮有力', '摇滚男中音，沙哑沧桑'],
        ['歌曲数', '9首（含序曲主题）', '7首（含终曲主题）'],
        ['工作量', '约35分钟', '约33分钟'],
        ['服装', '白色长衫整洁，束发，腰间佩剑', '白色长衫磨损发灰，散发，无剑'],
        ['关键词', '"我要飞"', '"我飞过了"'],
    ]
)

doc.add_heading('2.1.1 两人同台的时刻', level=3)

add_para('幕间《梦游天姥》——两人的第一次"相遇"', bold=True)
add_stage_direction('纱幕降下。年轻的李白(A)在纱幕后，老年的李白(B)在纱幕前。两人分别站在月亮灯箱两侧。')
add_para('这是李白(B)在全剧的第一次出场。他看到的第一个人——是二十年前的自己。两人不对话。年轻李白开口唱《大鹏赋》的原版旋律，老李白试图跟唱，嗓音已跟不上。他的声音沉入下方，变成和声的下声部。')

add_para('终曲《骑鲸》——两人的"最后一次告别"', bold=True)
add_stage_direction('所有角色依次登台。老李白站在中央。年轻李白从舞台后方缓步走来。两人面对面。年轻李白伸出手。老李白握住。两人一起走到月亮灯箱下——消失在月光中。')
add_para('音乐：《大鹏赋》主题由两人以相差八度的同旋律合唱。少年声部在上，沧桑声部在下。')

doc.add_heading('2.1.2 三版《大鹏赋》的声音设计', level=3)

add_table_with_data(
    ['出现位置', '演唱者', '声线', '编曲', '含义'],
    [
        ['序曲', '李白(A) 独唱', '清亮少年音', '竹笛+木吉他→全乐队', '"我要飞"'],
        ['将进酒结尾', '李白(A) + 全团', '摇滚高音嘶吼', '重金属+中国大鼓', '"我还在飞"'],
        ['终曲《骑鲸》', '李白(A)+李白(B) 合唱', '八度同旋律，少年在上沧桑在下', '管弦乐+合唱', '"飞本身即是意义"'],
    ]
)

add_quote('💡 这是单演员永远做不到的情感层次：唱出"我要飞"的那个少年，在终点等着老人归来。双演员把"一个人的一生"变成了"一个人与自己的对话"——而这恰好是李白诗歌中反复出现的母题（"举杯邀明月，对影成三人"）。')

doc.add_heading('2.2 女性角色（6位）', level=2)

add_table_with_data(
    ['角色', '声部', '身份', '独立弧光'],
    [
        ['许宛', '女中音', '第一任妻子', '她知道李白注定属于远方——仍选择给他一个家。这是成年人清醒的选择。'],
        ['玉真公主', '女高音', '玄宗之妹，入道女冠', '她代表了另一种自由——以退为进，在不自由中找到了自己的自由。'],
        ['杨玉环', '花腔女高音', '贵妃', '《霓裳》是她死前独唱。"倾国倾城"是标签——死亡面前她终于只为自己说话。'],
        ['宗倩', '次女高音', '最后一位妻子，修道者', '变卖家产救李白。不是痴情——"我就是这样选择活着的。"'],
        ['阿素', '女中音', '长安琵琶女（虚构）', '未经宫廷驯化的民间艺术的野性——她不需要李白来成就。'],
        ['月', '舞者（不唱）', '意象角色', '用身体承担全剧最抽象的追问——"人死后去了哪里？"'],
    ]
)

doc.add_heading('2.3 男性配角（6位）', level=2)

add_table_with_data(
    ['角色', '声部', '身份', '关键唱段'],
    [
        ['杜甫', '男中音', '诗人，最忠实的崇拜者', '《梦李白》'],
        ['唐玄宗', '男中音', '盛唐皇帝', '《沉香亭》'],
        ['高力士', '男低音', '宫廷权宦，体制的化身', '—'],
        ['李客', '男低音', '李白之父，西域商人', '《蜀道》（对唱）'],
        ['永王李璘', '男高音', '玄宗之子', '—'],
        ['当涂族叔', '男中音', '李白晚年的收留者', '—'],
    ]
)

add_quote('💡 角色均衡：6位女性 + 6位男性 + 月（非性别化意象）+ 李白双人 = 共14个主要表演者。女性戏份和歌曲数量与男性完全对等。')

# ── 三 ──
doc.add_heading('三、剧情架构（两幕制·120分钟）', level=1)

doc.add_heading('第一幕：大鹏（约55分钟）', level=2)
add_para('核心主题：一个天才如何相信自己可以飞', bold=True)

add_table_with_data(
    ['场次', '标题', '内容', '风格', '主要角色'],
    [
        ['序曲', '★《大鹏赋》', '全剧主题。少年李白离家', '交响摇滚', '李白(A)'],
        ['1', '《蜀道》', '少年在四川，父亲讲远方。出蜀', '民谣摇滚', '李白(A)、李客'],
        ['2', '《金陵酒肆》', '江南酒馆遇阿素。青春炽热', '轻摇滚+琵琶', '李白(A)、阿素'],
        ['3', '《终南山》', '隐居中遇玉真公主。论道论诗', '空灵二重唱', '李白(A)、玉真'],
        ['4', '《安陆》', '娶许宛，十年安陆。一个家', '温暖民谣', '许宛（独唱）'],
        ['5', '《长安·上》', '奉诏入京。长安繁华', 'Funk+中国风', '群像'],
        ['6', '《沉香亭》', '醉赋清平调。高力士脱靴', '宫廷摇滚', '李白(A)、杨玉环'],
        ['7', '★《将进酒》', '求去。告别长安。一幕高潮', '硬摇滚+大鼓', '李白(A)+全团'],
        ['幕间', '《梦游天姥》', '幻境。李白(A)与(B)首次相遇', '迷幻器乐+舞', '月、李白(A+B)'],
    ]
)

doc.add_heading('第二幕：明月（约50分钟）', level=2)
add_para('核心主题：飞过之后，如何降落', bold=True)

add_table_with_data(
    ['场次', '标题', '内容', '风格', '主要角色'],
    [
        ['8', '《梁园吟》', '遇宗倩。以诗为媒，以道为伴', '布鲁斯民谣', '李白(B)、宗倩'],
        ['9', '《醉眠秋共被》', '李白与杜甫相遇。文学史上最伟大的友谊', '吉他二重奏', '李白(B)、杜甫'],
        ['10', '《渔阳鼙鼓》', '安史之乱。杨玉环马嵬坡绝命', '重金属+破碎花腔', '杨玉环（独唱）'],
        ['11', '《夜郎》', '误投永王，以谋反罪流放。宗倩营救', '钢琴独白+二胡', '李白(B)、宗倩'],
        ['12', '《轻舟已过》', '遇赦东下。终于获得真正的自由', '公路摇滚', '李白(B)'],
        ['13', '《月下独酌》', '病困当涂。月夜追月入水', '古筝+吉他二重奏', '李白(B)、月'],
        ['终曲', '★《骑鲸》', '主题重现。两个李白合唱。全员送别', '管弦乐合唱', '全员'],
        ['谢幕', '《诗活着》', '脱古装，现现代装。诗还在', '流行摇滚', '全员'],
    ]
)

add_para('音乐总时长：约68分钟 | 对白转场：约45分钟 | 全剧总时长：约113分钟 ✓', bold=True)

# ── 四至七（略作精简，保留核心表格）──
doc.add_heading('四、主旋律设计：《大鹏赋》', level=1)

add_para('旋律：五声羽调式为基础，副歌融入♭7布鲁斯音。主歌6/8拍（大鹏振翅），副歌4/4摇滚。')
add_para('三次呈现的编曲演进：')
add_table_with_data(
    ['出现', '速度', '编曲', '演唱', '台词'],
    [
        ['序曲', '♩=72', '竹笛→木吉他→全乐队', '李白(A)独唱', '"大鹏一日同风起"'],
        ['一幕末', '♩=132', '战鼓+重金属+琵琶', '李白(A)+全团', '"大鹏不需要笼子"'],
        ['终曲', '♩=60', '海浪→大提琴→合唱', '李白(A+B)八度合唱', '"他骑上了鲸鱼"'],
    ]
)

doc.add_heading('五、歌曲清单（18首）', level=1)
add_table_with_data(
    ['#', '曲名', '演唱', '风格', '时长'],
    [
        ['1', '★《大鹏赋》', '李白(A)', '交响摇滚', '3\'30"'],
        ['2', '《蜀道》', '李白(A)+李客', '川江号子+民谣', '3\'00"'],
        ['3', '《金陵酒肆》', '李白(A)+阿素', '琵琶摇滚对唱', '3\'30"'],
        ['4', '《终南山》', '李白(A)+玉真', '空灵二重唱', '4\'00"'],
        ['5', '★《安陆》', '许宛', '温暖民谣', '3\'30"'],
        ['6', '《长安·上》', '群像', 'Funk+中国风', '3\'00"'],
        ['7', '★《清平调》', '李白(A)+杨玉环', '宫廷摇滚', '4\'30"'],
        ['8', '《脱靴》', '李白(A)+高力士', '讽刺摇滚', '2\'00"'],
        ['9', '★《将进酒》', '李白(A)+全团', '硬摇滚', '5\'00"'],
        ['—', '《梦游天姥》', '月+乐队', '迷幻器乐', '3\'00"'],
        ['10', '《梁园吟》', '李白(B)+宗倩', '布鲁斯民谣', '3\'30"'],
        ['11', '《梦李白》', '杜甫', '抒情吉他', '3\'00"'],
        ['12', '★《霓裳》', '杨玉环', '工业摇滚+花腔', '4\'00"'],
        ['13', '《千金赎》', '宗倩', '钢琴独白', '4\'00"'],
        ['14', '《我本楚狂人》', '李白(B)', '不插电Blues', '3\'00"'],
        ['15', '《轻舟已过》', '李白(B)', '公路摇滚', '3\'30"'],
        ['16', '《月下独酌》', '李白(B)+月', '古筝+吉他', '5\'00"'],
        ['17', '《骑鲸》', '全员', '管弦乐合唱', '4\'00"'],
        ['18', '《诗活着》', '全员', '流行摇滚', '3\'00"'],
    ]
)

doc.add_heading('六、音乐风格融合方案', level=1)
add_para('五大融合技法：')
add_para('① 琵琶 = 第二把电吉他（扫拂与失真同度齐奏）')
add_para('② 李白诗的吟诵节奏 = Hip-Hop（底鼓走trap节拍）')
add_para('③ 古筝摇指 = Shoegaze音墙（大量混响延迟处理）')
add_para('④ 京剧"叫板" = 摇滚"嘶吼"（李白愤怒时用花脸炸音）')
add_para('⑤ 五声旋律 + 爵士和声（II-V-I、三全音替代，古老皮肤+现代骨骼）')

doc.add_heading('七、女性角色的独立弧光', level=1)
add_para('六位女性各有一首独立唱段推动剧情。衡量标准：如果去掉李白这个角色，每一个女性角色仍然拥有完整的故事和情感逻辑。她们不是"谁的女人"——她们是"谁"。')

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 第二部分：校园落地方案
# ═══════════════════════════════════════════════════
add_para('第二部分：校园落地方案', 'Normal', bold=True, font_size=16, font_name='微软雅黑',
         color=RGBColor(0x1a, 0x1a, 0x2e), alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

doc.add_heading('八、校园舞台核心原则', level=1)
add_para('限制催生创造力。三个不减：歌不减、情不减、女性角色的独立弧光不减。')

doc.add_heading('九、舞美方案："一轮月亮 + 六扇屏风"', level=1)
add_para('全场唯一固定舞美。')
add_para('• 月亮灯箱：PVC管+LED灯带+白布，直径1.5m，材料费¥200-300')
add_para('• 六扇屏风：木框+白布双面，80cm×2m，四色（白/金/蓝/红），材料费¥500-800')
add_para('• 纱幕：3m×5m薄纱，仅幕间用，¥250')

doc.add_heading('十、乐队配置："五人成军"', level=1)
add_table_with_data(
    ['位置', '主奏', '兼任'],
    [
        ['键盘手（指挥）', '合成器工作站76键', '弦乐、编钟、钢琴、风琴音色'],
        ['吉他手', '电吉他+综合效果器', '—'],
        ['贝斯手', '电贝司', '—'],
        ['鼓手', '架子鼓', '中国大鼓、锣、钹'],
        ['国乐手', '琵琶', '兼古筝、笛子、二胡'],
    ]
)

doc.add_heading('十一、场景落地方案（逐场简化）', level=1)

scenes = [
    ('序曲：《大鹏赋》', '乐队全开。李白(A)从观众席通道走上台（打破第四面墙）。独站满月灯箱下，一束追光。屏风全白。'),
    ('第一场：《蜀道》', '两扇屏风翻绿面。6-8群演蹲伏→站起→手臂交织=蜀道山峰（肢体剧场）。竹笛前奏→民谣摇滚。'),
    ('第二场：《金陵酒肆》', '一张桌+两只碗+三把椅=酒肆。阿素做弹奏动作，国乐手在旁实奏琵琶——两人动作同步。'),
    ('第三场：《终南山》', '屏风全白。灰布铺地，群演在布下移动="云海"。两束蓝色追光分别打李白和玉真。成本最低效果最好的一场。'),
    ('第四场：《安陆》', '一盏暖色灯笼+两把椅。许宛独唱。全剧唯一完全不需要布景辅助的歌。越简单越动人。'),
    ('第五场：《长安·上》', '屏风全翻金面。群演16人手持金色纱巾流动。纱巾1元/条，比复杂服装有效。'),
    ('第六场：《沉香亭》', '屏风翻红面。群演蹲伏持红扇渐次站起=牡丹盛放。杨玉环从红屏风后走出。脱靴段灯光骤变刺眼白。'),
    ('第七场：《将进酒》', '鼓手走向中国大鼓。合唱团从观众席两侧边唱边走上来。李白站台口面向观众直接唱。结尾：全定格→灯灭→只留月亮。'),
    ('幕间：《梦游天姥》', '纱幕降下。月在纱幕后独舞。李白(B)纱幕前——第一次出场。李白(A)纱幕后唱主题。两人第一次"见面"。'),
    ('第八场：《梁园吟》', '一面屏风推台前。宗倩蘸水在屏风白布上写诗。灯光侧面打——水渍反光。字迹缓缓消失。'),
    ('第九场：《醉眠秋共被》', '两把椅+两个酒碗。乐器减到只剩键盘钢琴+国乐手轻弹。全剧最安静的一场。'),
    ('第十场：《渔阳鼙鼓》', '鼓手在中国大鼓上独奏30秒。群演把屏风一扇扇推倒。每声鼓倒一扇。杨玉环站废墟中独唱《霓裳》。'),
    ('第十一场：《夜郎》', '倒地屏风斜靠=牢笼。一束顶光从缝隙漏下。宗倩《千金赎》——钢琴+二胡。全剧最暗最催泪。'),
    ('第十二场：《轻舟已过》', '一面倒地屏风平放=船。群演在两侧持蓝布上下飘动=江水。全剧最明亮的段落。'),
    ('第十三场：《月下独酌》', '月亮灯箱满月→新月→微光。月从灯箱下走出第一次与李白共舞。入水：顶光→灭→3秒暗→月亮渐亮，李白已不在。'),
    ('终曲：《骑鲸》', '乐队+合唱团全上舞台。屏风全立，每扇不同色。所有角色依次上台。两个李白合唱主题。'),
    ('谢幕：《诗活着》', '演员逐一脱古装露出现代装。举起手机——屏幕上显示李白诗句。'),
]

for title, desc in scenes:
    doc.add_heading(title, level=2)
    add_para(desc)

doc.add_heading('十二、人员与预算', level=1)
add_para('核心团队：约25-30人（演员13人+乐队5人+群演/合唱8-12人+灯光1-2人+舞台监督1人+导演1人）')
add_para('预算：不含服装约¥1500-2000，含服装租赁约¥2500-5000。可申请学校社团经费。')

doc.add_heading('十三、排练时间线（8周）', level=1)
add_table_with_data(
    ['周次', '内容'],
    [
        ['1-2周', '围读剧本+音乐排练（分开）。录demo发演员自己练。'],
        ['3-4周', '分场景排练（不搭景，只走位+唱）。乐队+演员合音乐。'],
        ['5-6周', '联排。搭景。灯光编程。两个李白第一次在舞台上"相遇"。'],
        ['7周', '带妆联排。技术联排。解决所有问题。'],
        ['8周', '彩排 + 正式演出。'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 第三部分：完整舞台台本
# ═══════════════════════════════════════════════════
add_para('第三部分：完整舞台台本', 'Normal', bold=True, font_size=16, font_name='微软雅黑',
         color=RGBColor(0x1a, 0x1a, 0x2e), alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

doc.add_heading('十四、台本凡例', level=1)

add_para('本文使用以下格式约定：')
add_para('• 【】= 舞台提示（灯光、动作、走位、道具切换）')
add_para('• 角色名：= 对白')
add_para('• 《歌名》= 歌曲开始，歌词另起行缩进')
add_para('• → = 灯光/音乐转换节点')
add_para('• 台词中的引用诗句以" "标示')

add_para('本台本不是逐字逐句的精确剧本——对白部分是一个"情感和信息的路线图"。导演和演员应在排练中根据实际效果调整具体措辞。但每场戏的结构、人物的情感转折、关键台词的核心信息不应偏离。', italic=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 第一幕台本
# ═══════════════════════════════════════════════════
doc.add_heading('十五、第一幕台本：大鹏', level=1)

# ── 序曲 ──
doc.add_heading('序曲：《大鹏赋》', level=2)
add_para('时长：约4分钟 | 演唱：李白(A) | 风格：交响摇滚', bold=True)

add_stage_direction('剧场灯光全暗。观众席后排的门打开——一束光从门外射入。年轻的李白(A)从观众席通道大步走向舞台。他穿过观众，有时与某人对视。他白衫整洁，腰间佩剑。')
add_stage_direction('走到台口时，月亮灯箱缓缓亮起——满月。六扇屏风全是白色面。追光跟住他。')
add_stage_direction('他站定。转身面对观众。钢琴（键盘）开始——简单的下行音阶，重复四次。竹笛（国乐手）进入——只有一个音，拖了八拍长。')

add_dialogue('李白(A)', '（对着观众，不像是表演，更像是聊天）我五岁的时候，我父亲跟我说了一句话——他说：思辰……不对，太白。他说：太白，这世上有些人，生来就是为了抵达。有些人呢，生来就是为了出发。他说：你猜你是哪一种？')

add_stage_direction('吉他、贝斯、鼓同时进入。全编制爆炸。主题旋律起。李白(A)抬起头。20岁，眼睛里全是光。')

add_dialogue('李白(A)', '我二十岁了。我不知道答案。但我决定——先出发再说。')

add_stage_direction('唱《大鹏赋》。')

# 歌词用缩进楷体
add_lyric_line('【第一段】', 0.5)
add_lyric_line('大鹏一日同风起——')
add_lyric_line('扶摇直上九万里。')
add_lyric_line('假令风歇时下来——')
add_lyric_line('犹能簸却沧溟水。')

add_stage_direction('副歌——乐队全开。李白在舞台前端，追光缩小到只照亮他上半身。')

add_lyric_line('【副歌】', 0.5)
add_lyric_line('我不是鸟——我是风本身！')
add_lyric_line('不需要方向——因为所有方向都是我的。')
add_lyric_line('我叫李白——我不需要向你解释。')
add_lyric_line('大鹏飞过的地方——就是天空。')

add_stage_direction('尾声——月亮灯箱亮度推到最大。李白转身，走向舞台深处。追光收。屏风在他经过时一扇一扇被群演推到舞台后区。他消失在屏风后面。')
add_stage_direction('灯光渐灭。只留月亮。')

doc.add_paragraph()

# ── 第一场：蜀道 ──
doc.add_heading('第一场：《蜀道》', level=2)
add_para('时长：约7分钟 | 演唱：李白(A)+李客 | 风格：川江号子+民谣摇滚', bold=True)

add_stage_direction('屏风两扇翻深绿面。月亮灯箱亮度减半。群演8人穿深色衣，在舞台后区以肢体构成"山"的意象——蹲伏、缓慢站起、手臂交叠、身体倾斜。')
add_stage_direction('李客（父亲，约45岁）站在舞台右侧。他穿着胡服，带有明显的西域风格。李白(A)背着简单的行囊，站在左侧。')

add_dialogue('李客', '四川。你出生的地方。你五岁的时候，我带你翻过那座山——（指向群演组成的"山"）你还记不记得？')
add_dialogue('李白(A)', '（摇头）我只记得……云。')
add_dialogue('李客', '（笑）云？')
add_dialogue('李白(A)', '雾从山底往上飘，像倒着下的雨。我以为翻过那座山，就是天边了。（停顿）父亲，天边到底在哪里？')
add_dialogue('李客', '（长久沉默。然后从怀里掏出一卷小小的羊皮地图）我年轻的时候，从碎叶城出发，走过龟兹、于阗、敦煌——进到玉门关的时候，我以为我到了。后来才发现，那只是开始。')

add_stage_direction('他展开地图。地图上没有任何字，只有几条墨线——河流、山脉。')

add_dialogue('李客', '太白，你要记住——你身上流的不是只有中原的血。你的祖母是胡人。你的眼睛看得比他们远，不是因为你站得高——是因为你同时属于两个世界。')
add_dialogue('李白(A)', '所以我才没办法在一个地方停下来？')
add_dialogue('李客', '不。你能停下来。你只是不需要停下来。')

add_stage_direction('竹笛（国乐手）吹出蜀地民歌的旋律片段。吉他进入，分解和弦。')

add_dialogue('李白(A)', '父亲。我要走了。')
add_dialogue('李客', '我知道。')

add_stage_direction('唱《蜀道》。')

add_lyric_line('【李白(A) 主歌】', 0.5)
add_lyric_line('蜀道难——难不过一个年轻的膝盖。')
add_lyric_line('山多高——高不过一双想看遍天下的眼睛。')
add_lyric_line('父亲说：去吧，把地图像我一样——')
add_lyric_line('画满。画到没有空白的地方为止。')

add_lyric_line('【李客 第二段】', 0.5)
add_lyric_line('我走过的地方——留在你的脚下。')
add_lyric_line('我没有走完的——你可以接着走。')
add_lyric_line('只是记住：风沙大的时候——')
add_lyric_line('低头看看影子——那是家的方向。')

add_lyric_line('【父子对唱副歌】', 0.5)
add_lyric_line('蜀道难，难于上青天——')
add_lyric_line('可青天之上，是更多的路。')
add_lyric_line('蜀道难，难于上青天——')
add_lyric_line('可你的名字叫太白——太白是星星的名字。')

add_stage_direction('歌曲结束。李白(A)向父亲深鞠一躬。李客把羊皮地图塞进他手里。转身走入屏风后，不再回头。')
add_stage_direction('李白(A)把地图揣入怀中，深吸一口气，迈出第一步。群演组成的"山"在他经过时慢慢散开——山为他让路。')
add_stage_direction('灯光切为明亮的金色面光。转场。')

doc.add_paragraph()

# ── 第二场：金陵酒肆 ──
doc.add_heading('第二场：《金陵酒肆》', level=2)
add_para('时长：约8分钟 | 演唱：李白(A)+阿素 | 风格：轻摇滚+琵琶即兴solo', bold=True)

add_stage_direction('一张方桌，两只酒碗，一把琵琶靠在椅边。灯光是温暖的金黄色。江南。')
add_stage_direction('阿素——约23岁，穿着利落的短衣，正在调琵琶弦。她是教坊最好的琵琶手，但她身上没有一丝宫廷气——她是市井的、野生的、鲜活的。')
add_stage_direction('李白(A)推门进来（实际上只是从舞台侧幕走出，但他的身体语言告诉你——他推开了一扇门）。他比上一场多了些尘土——已经在路上走了几个月。')

add_dialogue('阿素', '（头也不抬，继续调弦）打酒还是住店？')
add_dialogue('李白(A)', '听琴。')
add_dialogue('阿素', '（终于抬头，打量他）听琴？你懂琵琶？')
add_dialogue('李白(A)', '不懂。但我知道好听的琴声会让酒变甜。（晃了晃酒壶）我试过好几次了——屡试不爽。')
add_dialogue('阿素', '（笑出了声）你是诗人？')
add_dialogue('李白(A)', '你怎么知道？')
add_dialogue('阿素', '只有诗人会这样说话——把一句话绕三个弯，就为了说"你的琵琶弹得不错"。')

add_stage_direction('阿素把琵琶抱起来。随手拨了一个音——清澈，穿透。')

add_dialogue('阿素', '你说酒会变甜？那我们做个交易。你给我一首诗——现写。我满意了，请你一壶酒。不满意——你请我一壶。')
add_dialogue('李白(A)', '（毫不犹豫）成交。')

add_stage_direction('阿素开始弹琵琶。即兴的旋律——轻快、俏皮、带着挑衅。国乐手在舞台侧前方实际演奏，阿素演员的动作与国乐手完全同步。')

add_stage_direction('唱《金陵酒肆》。')

add_lyric_line('【阿素 主歌】', 0.5)
add_lyric_line('风吹柳花满店香——')
add_lyric_line('吴姬压酒唤客尝。')
add_lyric_line('这位公子——你从哪里来？')
add_lyric_line('一身的风沙像是走了很远。')

add_lyric_line('【李白(A) 主歌】', 0.5)
add_lyric_line('我从蜀中来——沿着江水往东走。')
add_lyric_line('每到一个地方就问——天边还有多远？')
add_lyric_line('所有人都在笑我——')
add_lyric_line('可是——你的琴声好像没有笑。')

add_lyric_line('【阿素 副歌前】', 0.5)
add_lyric_line('因为我的琴也想去天边——')
add_lyric_line('但它只有四根弦，飞不出去。')

add_lyric_line('【李白(A) 回应】', 0.5)
add_lyric_line('那就给我听——我帮你飞。')
add_lyric_line('用我的诗当翅膀——够不够？')

add_lyric_line('【合唱副歌】', 0.5)
add_lyric_line('金陵的酒——不醉人，醉的是今晚的琴声。')
add_lyric_line('金陵的酒——不醉人，醉的是少年人的约定。')
add_lyric_line('来日你若路过长安——')
add_lyric_line('记得这壶酒——和这首歌。')

add_stage_direction('琵琶solo——24小节。国乐手走到舞台中央。阿素站在她身旁，手搭在她肩上。两人的身体同频呼吸。这是全剧第一次让"音乐本身"成为被凝视的焦点——不服务于剧情，音乐就是剧情。')

add_stage_direction('solo结束。李白(A)往桌上扔了一小块银子。')

add_dialogue('李白(A)', '酒钱。')
add_dialogue('阿素', '诗还没写呢。')
add_dialogue('李白(A)', '（边走边回头，笑）写过了。刚才那首歌——就是诗。')

add_stage_direction('李白(A)下。阿素独自站在桌前，拿起他喝过的酒碗，看了一眼。轻轻摇头，自己喝了一口。')
add_stage_direction('转场。灯光从金黄渐变为冷白。')

doc.add_paragraph()

# ── 第三场：终南山 ──
doc.add_heading('第三场：《终南山》', level=2)
add_para('时长：约9分钟 | 演唱：李白(A)+玉真公主 | 风格：New Age+古筝，空灵二重唱', bold=True)

add_stage_direction('灯光：月光蓝。屏风全部转白色面。一块灰布铺在舞台地面，8个群演蹲在布下缓慢移动——"云海"。')
add_stage_direction('舞台后方，玉真公主坐在一块石头（道具箱蒙灰布）上。她穿着素雅的道袍——不是宫中公主的华丽，是山间修行者的简朴。但她的气质仍然带着不可忽视的贵气。')
add_stage_direction('李白(A)从舞台左侧上。他的衣服比上一场更旧了，但精神更好。他看见玉真公主——愣了一下。')

add_dialogue('李白(A)', '（自语般）终南山的云……什么时候变成人了？')
add_dialogue('玉真公主', '（没有转头）你就是那个在金陵用一首诗换了壶酒的李白？')
add_dialogue('李白(A)', '消息传得这么快？')
add_dialogue('玉真公主', '我听到的不是"消息"。是她的琵琶。阿素把你们那天的曲子记了下来——教坊的女孩们都在弹。（终于转头看他）你知道长安城现在最流行的曲子，是一个从四川来的无名诗人写的吗？')

add_stage_direction('李白(A)走过去，在离她几步远的地方坐下。两人之间的空间——保持距离，但有一种奇怪的熟悉感。')

add_dialogue('李白(A)', '所以你是……？')
add_dialogue('玉真公主', '一个住在山上的人。和你一样——（顿了顿）也是一个不想回家的人。')
add_dialogue('李白(A)', '家有什么不好？')
add_dialogue('玉真公主', '家没有什么不好。是"回家"这两个字——（她抬头看月亮灯箱）——每次听到，都觉得是在说"回到别人给你安排好的位置上去"。')

add_stage_direction('长时间沉默。古筝（国乐手）开始弹——缓慢的、空灵的琶音。')

add_dialogue('李白(A)', '所以你住到了山上。')
add_dialogue('玉真公主', '所以我住到了山上。修道是一个很好的理由——没有人会追问。你呢？为什么一直在路上？')
add_dialogue('李白(A)', '（想了一下）因为我怕——如果停下来，就会被装进一个叫"李白"的瓶子里。然后所有人都说：看，这是李白。而瓶子里的人，已经死了。')

add_stage_direction('玉真公主第一次认真地看他。不是打量——是认出。')

add_dialogue('玉真公主', '我可以帮你。')
add_dialogue('李白(A)', '帮我什么？')
add_dialogue('玉真公主', '去长安。去皇帝面前。你可以站到更高的地方——不是因为那个地方好，而是因为……（她站起来）大鹏不应该只在江南的小酒馆里飞。')

add_stage_direction('唱《终南山》。')

add_lyric_line('【玉真公主 第一段】', 0.5)
add_lyric_line('终南山的云——不需要方向。')
add_lyric_line('它们只是飘着——就已经足够了。')
add_lyric_line('可是你不是云——你是风。')
add_lyric_line('风如果不吹——它就不再是风了。')

add_lyric_line('【李白(A) 回应】', 0.5)
add_lyric_line('长安远吗——比蜀道更远？')
add_lyric_line('长安高吗——比终南山更高？')
add_lyric_line('如果我去了——如果我留下来了——')
add_lyric_line('我还会不会是我自己？')

add_lyric_line('【玉真公主 第二段】', 0.5)
add_lyric_line('我选择了"不"——不嫁人。不入宫。不参与天下的棋局。')
add_lyric_line('这是我的自由——它很小，但它是我自己的。')
add_lyric_line('你的自由不一样——你需要去。去了，再回来——')
add_lyric_line('你才会知道：一直不去，也是不自由。')

add_lyric_line('【二重唱 副歌】', 0.5)
add_lyric_line('（李）我不属于任何地方——有人告诉我这是错的。')
add_lyric_line('（玉）不属于任何地方——也许这意味着你属于所有地方。')
add_lyric_line('（李）我怕长安会改变我。')
add_lyric_line('（玉）你怕的不是改变——你怕的是变成你不认识的那个人。')
add_lyric_line('（合）但认识那个不认识的人——也许正是这趟路的意义。')

add_stage_direction('歌曲结束。玉真公主从袖中取出一封信，递给李白(A)。')

add_dialogue('玉真公主', '这是给皇兄的举荐信。但到了长安，要不要递上去——你自己决定。')
add_dialogue('李白(A)', '（接过信）如果我不递呢？')
add_dialogue('玉真公主', '（微微一笑，转身向山深处走去）那你至少知道——长安是有的。有一个需要你去的地方，和有一个需要你回去的地方——是两件完全不同的事。')

add_stage_direction('她消失在屏风后。李白(A)独自留在"云海"中。他看着手里的信。折好。放进怀里——和父亲的羊皮地图放在一起。')
add_stage_direction('灯光渐暗。转场。')

doc.add_paragraph()

# ── 第四场：安陆 ──
doc.add_heading('第四场：《安陆》', level=2)
add_para('时长：约8分钟 | 演唱：许宛（独唱）| 风格：温暖民谣，钢琴+二胡', bold=True)

add_stage_direction('灯光：温暖的琥珀色。六扇屏风中两扇翻为淡黄色面（可以用暖光打在白色面上代替）。一盏灯笼挂在舞台中央上方（实际是手持或悬挂的小道具灯笼）。')
add_stage_direction('许宛——约28岁，女中音。穿着素雅的居家服装，不是华丽的贵族女子，而是安静、笃定、有力量的女人。她坐在椅子上，手中做针线。')
add_stage_direction('李白(A)站在舞台一侧，靠着屏风，看着她。他的服装换成了稍微体面一些的——入赘许家后的打扮。但他腰间仍然佩着剑。')

add_dialogue('许宛', '（头也不抬，语气平和）今天又没写？')
add_dialogue('李白(A)', '写了。写了两句，然后……看院子里的桃花去了。')
add_dialogue('许宛', '桃花开得是很好。')
add_dialogue('李白(A)', '（突然）宛娘。')
add_dialogue('许宛', '嗯？')
add_dialogue('李白(A)', '如果有一天……我走了。你会恨我吗？')

add_stage_direction('许宛的手停了一下。但只是一下。然后继续缝。')

add_dialogue('许宛', '太白。我嫁给你的时候，就知道你不会在一个地方待一辈子。')
add_dialogue('李白(A)', '那你为什么还……')
add_dialogue('许宛', '（放下针线，抬头看他）因为我想要一个家——哪怕只有十年。十年的家也是家。十年之后你走了——它还是家。只不过成了一个人的家。')

add_stage_direction('长时间的沉默。远处似乎有鸟叫。')

add_dialogue('许宛', '你不用觉得欠我。我选择你——不是因为我不知道你会走。恰恰是因为我知道。（她站起来，走到他面前）一个不知道远方是什么的人，他留在我身边——他的心也不在这里。但你不一样。你看过蜀道和金陵——你选择在这里停下来——这说明这一刻，你是真的在这里。真的在——就够了。')

add_stage_direction('她伸手，把他腰间的剑轻轻按了一下——像是把一件随时会飞走的东西暂时压住。')

add_dialogue('许宛', '剑别解。等你想走的时候——再解。')
add_dialogue('李白(A)', '（声音在喉咙里）宛娘……')
add_dialogue('许宛', '（回到椅子上，重新拿起针线）去写你的诗吧。桃花明天就谢了。')

add_stage_direction('唱《安陆》。独唱——许宛一个人坐在椅子上。键盘手切到钢琴音色。国乐手拉二胡。全剧最温暖的段落。')

add_lyric_line('【主歌】', 0.5)
add_lyric_line('安陆的春天——桃花开了满院子。')
add_lyric_line('你说要给我写一首诗——')
add_lyric_line('写到花谢了——写到冬来了——')
add_lyric_line('你还在想第一句。')

add_lyric_line('这没有关系——我不需要你的诗。')
add_lyric_line('我需要的——是你在想第一句的时候——')
add_lyric_line('眼睛看着院子里的桃花——')
add_lyric_line('而桃花也在看你。')

add_lyric_line('【副歌】', 0.5)
add_lyric_line('你是风——我知道你是风。')
add_lyric_line('风不会停下来——停下来就不是风了。')
add_lyric_line('可风在吹过桃花的那个春天——')
add_lyric_line('那一刻的风——是属于桃花的。')

add_lyric_line('【第二段主歌】', 0.5)
add_lyric_line('他们说你应该找个安分的人——')
add_lyric_line('找一个不会在半夜忽然坐起来——')
add_lyric_line('对着月亮发呆的人。')
add_lyric_line('可是——安分的人，他看桃花的时候——')
add_lyric_line('桃花只是桃花。而你不是。')

add_lyric_line('【桥段】', 0.5)
add_lyric_line('十年——够了。')
add_lyric_line('一个家——够了。')
add_lyric_line('等你走到走不动的那天——')
add_lyric_line('如果你还记得回来的路——')
add_lyric_line('桃花还开着。我还在这里。')

add_lyric_line('【副歌 重复渐弱】', 0.5)
add_lyric_line('你是风——我知道你是风。')
add_lyric_line('风不需要为任何人停下来。')
add_lyric_line('可是风吹过的那棵桃树——')
add_lyric_line('永远记得风的方向。')

add_stage_direction('歌曲结束时，二胡的最后一个长音拖了很久——直到消失。许宛把手中缝好的东西展开——是一件新的白色长衫。她把它叠好，放在椅子上。起身。走进屏风后。不再回头。')
add_stage_direction('李白(A)走到椅子前，拿起那件长衫。抱在怀里。背对观众站了很久。')
add_stage_direction('灯光缓慢收至只剩下月亮灯箱。转场。')

doc.add_paragraph()

# ── 第五场：长安·上 ──
doc.add_heading('第五场：《长安·上》', level=2)
add_para('时长：约8分钟 | 演唱：群像 | 风格：Funk+琵琶，中国风舞曲', bold=True)

add_stage_direction('灯光转变：金色面光+金色顶光同时打亮——全剧最明亮的场景之一。六扇屏风全部翻到金色面。')
add_stage_direction('群演16人从舞台两侧涌入，每人手持一条金色纱巾。纱巾在灯光下像流水一样飘动——这是长安，这是盛唐的巅峰。')
add_stage_direction('李白(A)穿着许宛缝制的白色长衫站在舞台中央。他周围的人流越走越快——他是漩涡中心静止的一点。')

add_dialogue('李白(A)', '（对着观众，语气从震惊到眩晕）长安。我到了。我到了——可是我找不到自己的脚了。每个人都在走，每个人都在笑，每个人都在说——"李白是谁？"（苦笑）我以为全天下都该知道我是谁。但长安有太多人了——多到一个人的名字掉进去，就像一滴水掉进黄河。')

add_stage_direction('群舞——Funky的节奏，琵琶做切分。皇帝不上场——长安本身就是一个角色。')

add_dialogue('李白(A)', '（在歌曲中间插入，对观众）玉真公主的信在我怀里。三天了——我没有递。我在旅馆里躺着，看天花板上的裂纹。我想——如果我把信递了，皇帝看了，他让我做官——我该怎么办？每天早上去上班？（笑出声）钟声一响，李白夹着一个公文包——（摇头）不行。不行不行不行。')

add_stage_direction('群像唱《长安·上》。')

add_lyric_line('【群演合唱 主歌】', 0.5)
add_lyric_line('长安长安——天下之中！')
add_lyric_line('东西两市——胡姬酒肆——波斯琉璃——西域香料——')
add_lyric_line('左脚踩着节拍，右脚踩着平康坊的石板路。')
add_lyric_line('每一个来长安的人——都想成为长安。')

add_lyric_line('【李白(A) 插入】', 0.5)
add_lyric_line('（在群演喧哗声中，对着观众）')
add_lyric_line('我闻到烤羊肉的味道了——不对，那是檀香。')
add_lyric_line('我听见琵琶了——不对，那是——那是整个长安在唱歌。')

add_lyric_line('【群演副歌】', 0.5)
add_lyric_line('来啊！把最好的酒端上来！')
add_lyric_line('来啊！把最亮的灯点起来！')
add_lyric_line('长安没有夜晚——只有一场永远不会散的宴席！')

add_stage_direction('歌曲在高潮戛然而止。群演瞬间定格——所有人保持金色纱巾飘在空中的姿势。灯光从金色骤变为冷白。')
add_stage_direction('一个低沉的声音从舞台后区传来。')

add_dialogue('李白(A)', '（在凝固的人群中，独自走动）然后我明白了。长安不只是一座城——它是一面镜子。你在它面前，照出来的不是你是谁——而是你愿意变成谁。（停住）我不愿意变成任何别的人。')

add_stage_direction('他摸了摸怀里的信。')
add_dialogue('李白(A)', '但我还是递了。因为——大鹏需要飞过长安。哪怕只是为了跟长安说一声——"我来了。我看见了。再见。"')

add_stage_direction('定格解除。群演如潮水般退下。灯光转换。')

doc.add_paragraph()

# ── 第六场：沉香亭 ──
doc.add_heading('第六场：《沉香亭》', level=2)
add_para('时长：约10分钟 | 演唱：李白(A)+杨玉环+唐玄宗+高力士 | 风格：华丽宫廷摇滚', bold=True)

add_stage_direction('灯光：金色顶光+红色侧光。屏风全部翻红色面。群演蹲伏在屏风前，每人手持一把红扇——他们缓慢站起、展开扇子，像牡丹绽开。')
add_stage_direction('杨玉环——25岁左右，花腔女高音。她穿着正红色的华丽长裙。她从红色屏风中缓步走出——她不是走进这个场景，她是这个场景的源头。')
add_stage_direction('唐玄宗——50岁左右，男中音。穿着便服而非朝服，显出他此刻只是一个赏花的男人。他坐在舞台后区的一把椅子上。')
add_stage_direction('高力士——男低音，站在玄宗身侧，永远挺直。')

add_stage_direction('李白(A)从舞台左侧上。他穿着一身新衣服——宫廷赐的翰林服。但他的身体语言告诉你他不舒服——像一个穿着别人衣服的人。')

add_dialogue('高力士', '（声音冷而平）李白，陛下宣你。今日沉香亭牡丹盛开，命你赋诗。')
add_dialogue('李白(A)', '（看了高力士一眼）赋诗需要酒。')
add_dialogue('高力士', '你已经醉了。')
add_dialogue('李白(A)', '醉和不醉——对诗来说不是两个状态。醉是墨，不醉是白纸。你要白纸，还是诗？（不等高力士回答）我知道——诗。那就再给我一壶。')

add_stage_direction('玄宗轻轻笑了一声。高力士不再说话。有人递给李白一壶酒。李白仰头灌了一口。')

add_dialogue('唐玄宗', '（温和地）李白，你看这牡丹——比去年开得更好吗？')
add_dialogue('李白(A)', '（看向牡丹，又看向杨玉环）牡丹不看去年。牡丹只管今年——今年开了，就拼尽全力地开。它不比较。（顿了顿，迎上玄宗的目光）就像陛下身边的这位——（他微微向杨玉环的方向欠身）——不需要和任何人比较。')

add_stage_direction('杨玉环第一次开口。她的声音不高，但在整个舞台上都听得见。')

add_dialogue('杨玉环', '你叫李白。（不是疑问句）教坊的女孩们都在弹你的曲子。')
add_dialogue('李白(A)', '（有点意外）贵妃也知道？')
add_dialogue('杨玉环', '我什么都听。在宫里——听是唯一不需要得到批准的事。')

add_stage_direction('李白(A)与杨玉环对视了一瞬间。两个人都明白了——他们是同类。在金色的笼子里，他们都是那个在半夜对着月亮发呆的人。')

add_dialogue('唐玄宗', '（没有注意到这个对视，或者假装没有）李白，你先赋诗。贵妃会唱。')
add_dialogue('李白(A)', '（把酒壶放在地上，站直）遵旨。')

add_stage_direction('唱《清平调》。')

add_lyric_line('【李白(A) 第一段——对着牡丹，但每个字都是给杨玉环的】', 0.5)
add_lyric_line('云想衣裳花想容——')
add_lyric_line('春风拂槛露华浓。')
add_lyric_line('若非群玉山头见——')
add_lyric_line('会向瑶台月下逢。')

add_lyric_line('【杨玉环 回应——花腔，华丽但带一丝冷】', 0.5)
add_lyric_line('一枝红艳露凝香——')
add_lyric_line('云雨巫山枉断肠。')
add_lyric_line('借问汉宫谁得似——')
add_lyric_line('可怜飞燕倚新妆。')

add_lyric_line('【李白(A) 第三段——比前两段更慢，更轻】', 0.5)
add_lyric_line('名花倾国两相欢——')
add_lyric_line('长得君王带笑看。')
add_lyric_line('解释春风无限恨——')
add_lyric_line('沉香亭北倚阑干。')

add_lyric_line('【杨玉环 副歌——花腔，在"无限恨"三个字上陡然拔高然后坠落】', 0.5)
add_lyric_line('无限恨——春风能解释什么？')
add_lyric_line('春风自己——也是被风吹来的。')
add_lyric_line('沉香亭的牡丹开不过这一季——')
add_lyric_line('可是亭子会一直在。亭子永远在。')

add_stage_direction('歌曲到此处——李白(A)忽然把笔扔在地上。声音在静默中很响。')

add_dialogue('李白(A)', '（醉意上涌，但每个字都很清楚）高将军——我脚上的靴子有点紧。能帮我脱一下吗？')

add_stage_direction('全场冻结。玄宗脸上的笑意凝固。杨玉环的眼神从李白移到高力士——她微微皱了一下眉，像是担心，又像是在忍住笑意。')
add_stage_direction('高力士一动不动。灯光逐渐变成刺眼的白——舞台变成审讯室。')

add_dialogue('高力士', '（非常缓慢地）……你说什么？')
add_dialogue('李白(A)', '（没有退缩，但语气中没有恶意——他确实醉了，但醉得清醒）我说——靴子。紧。能不能——帮我脱一下。')

add_stage_direction('长时间的沉默。玄宗没有开口。杨玉环转过了脸。')

add_dialogue('高力士', '（声音如冰）李白。你记住今天的沉香亭。')

add_stage_direction('他跪下。伸手。李白(A)低头看着这个帝国最有权势的宦官跪在自己面前——他的表情不是得意，而是一种巨大的荒诞感。')

add_stage_direction('高力士的手触到靴子的瞬间——灯光全变红。音乐——一段不和谐的弦乐（键盘手）+ 低沉的贝斯线。')

add_stage_direction('唱《脱靴》——短小的讽刺摇滚。')

add_lyric_line('【李白(A) 独唱】', 0.5)
add_lyric_line('一双靴子——脱下来只需要三秒。')
add_lyric_line('穿上——却需要一辈子。')
add_lyric_line('高将军的手很稳——')
add_lyric_line('稳到可以扶起一个帝国——也可以掐断一个人的喉咙。')

add_lyric_line('我脱掉的是靴子——')
add_lyric_line('他脱掉的是什么——')
add_lyric_line('没有人敢问。')

add_stage_direction('靴子脱下。李白(A)赤脚站在舞台中央。高力士站起来——他的脸上没有任何表情。转身。走回玄宗身侧。')
add_stage_direction('玄宗站起来。他看了李白(A)一眼——表情复杂。转身。杨玉环跟在他身后——但在经过李白时，她停了一瞬间。')

add_dialogue('杨玉环', '（只动嘴唇，几乎不出声）走。')

add_stage_direction('玄宗、高力士、杨玉环依次走入屏风后。群演收扇，退场。李白(A)独自站在红色屏风前。赤脚。灯光从红色缓慢变成月光蓝。他抬起头——月亮灯箱还在那里。')

add_dialogue('李白(A)', '（轻声）许宛是对的。我需要来长安——不是为了留下来。是为了知道——这里不是我的家。')

add_stage_direction('他捡起地上的酒壶，仰头喝完最后一口。把空壶放在舞台上。赤脚走向舞台前方——直接进入《将进酒》。')

doc.add_paragraph()

# ── 第七场：将进酒 ──
doc.add_heading('第七场：《将进酒》', level=2)
add_para('时长：约8分钟 | 演唱：李白(A)+全团 | 风格：硬摇滚+中国大鼓。第一幕高潮。', bold=True)

add_stage_direction('灯光仍然以月光蓝为底，但鼓手已从架子鼓位走到舞台后区的中国大鼓前。合唱团8-12人从观众席两侧通道缓步走上舞台——他们边走边哼着《大鹏赋》的无词旋律。')
add_stage_direction('李白(A)赤脚站在台口，面向观众。他把翰林服的外袍脱下来——扔在一边。露出里面许宛缝的那件白色长衫。')

add_dialogue('李白(A)', '（对观众，语气平静，但底下全是火）很多人问我——李白，你为什么离开长安？皇帝对你不错。高力士的事——是你自己找的。你为什么就不能——忍一忍？')

add_dialogue('李白(A)', '（停顿。然后——几乎是自言自语）因为忍一忍，就变成了另一个人。而这世上——不缺一个能忍的李白。缺的只有一个——敢走的李白。')

add_stage_direction('中国大鼓——第一下。沉闷的，像心跳。吉他失真音色——一个长音，从低声部慢慢爬升。')

add_dialogue('李白(A)', '今晚——所有来送我的朋友——所有在长安认识的、不认识的——（他转过身，对着舞台上逐渐聚集的合唱团和群演）这杯酒——不敬天子。不敬贵妃。不敬长安。')

add_stage_direction('鼓——第二下。更响。贝斯进入。')

add_dialogue('李白(A)', '（把酒壶高高举起，对着月亮灯箱）敬月亮。敬酒。敬还没走的这条路。')

add_stage_direction('鼓——第三下。乐队全编制爆炸进入。')

add_para('── 《将进酒》歌曲开始 ──', bold=True)

add_stage_direction('【引子】电吉他长音 + 中国大鼓三声重击 + 琵琶扫弦引入。速度 ♩=132。')

add_lyric_line('【主歌 A1 — 李白(A) 独唱，乐队轻伴奏】', 0.5)
add_lyric_line('君不见——黄河之水天上来——')
add_lyric_line('奔流到海不复回。')
add_lyric_line('君不见——高堂明镜悲白发——')
add_lyric_line('朝如青丝暮成雪。')

add_stage_direction('吉他失真进入——功率和弦。鼓手从中国鼓回到架子鼓，打出硬摇滚节奏。')

add_lyric_line('【主歌 A2 — 李白(A) ，声压渐强】', 0.5)
add_lyric_line('人生得意须尽欢——莫使金樽空对月！')
add_lyric_line('天生我材必有用——千金散尽还复来！')
add_lyric_line('烹羊宰牛且为乐——会须一饮三百杯！')

add_stage_direction('合唱团加入——和声叠在李白的主旋律上方。')

add_lyric_line('【合唱团进入 — 重复"三百杯"三遍，渐强】', 0.5)
add_lyric_line('三百杯！三百杯！三——百——杯！')

add_stage_direction('乐队骤停。只留架子鼓的hi-hat——踩镲的嘶嘶声。全场灯光压暗到只剩追光打在李白身上。')

add_lyric_line('【桥段 — 李白(A) 近似吟诵，鼓走trap groove】', 0.5)
add_lyric_line('岑夫子——丹丘生——将进酒——杯莫停。')
add_lyric_line('与君歌一曲——请君为我倾耳听。')

add_stage_direction('全编制回归。追光扩至整个舞台。金色+红色顶光同时亮。全体演员站到台前——打破第四面墙。')

add_lyric_line('【大合唱 副歌 — 李白(A) 主唱，全员和声】', 0.5)
add_lyric_line('钟鼓馔玉不足贵——但愿长醉不愿醒！')
add_lyric_line('古来圣贤皆寂寞——惟有饮者留其名！')
add_lyric_line('陈王昔时宴平乐——斗酒十千恣欢谑！')
add_lyric_line('主人何为言少钱——径须沽取对君酌！')

add_stage_direction('间奏——电吉他solo 16小节 + 琵琶同度齐奏（全剧标志性的"古今碰撞"声音）。鼓手走中国大鼓和架子鼓交替的复合节奏。')

add_lyric_line('【李白(A) 独唱 — 乐队降到只剩贝斯+鼓】', 0.5)
add_lyric_line('五花马——千金裘——')
add_lyric_line('呼儿将出换美酒——')
add_lyric_line('与尔同销——万——古——愁！')

add_stage_direction('"万——古——愁"三个字每字之间隔一拍。每一下——乐队全奏一个和弦。鼓——重击。万！（咚）古！（咚）愁！（咚——）')

add_stage_direction('最后一遍副歌——全编制+合唱团+所有演员。速度推到最快。金色灯光全开——整座舞台像一个着了火的宫殿。')

add_lyric_line('【终副歌 — 全团，速度 ♩=144】', 0.5)
add_lyric_line('将进酒——杯莫停——！')
add_lyric_line('与君歌一曲——请君为我倾耳听——！')
add_lyric_line('钟鼓馔玉不足贵——但愿长醉不愿醒——！')
add_lyric_line('古来圣贤皆寂寞——惟有饮者留其名——！')

add_stage_direction('歌曲最后三个和弦——全体演员定格——手臂高举——嘴张着——像被凝固在一声呐喊的顶点。')
add_stage_direction('灯光——骤灭。')
add_stage_direction('黑暗中，只有月亮灯箱还亮着。三秒后，月亮也缓慢调暗——但不是全灭，而是留着一层微光。')

add_dialogue('李白(A)', '（在完全的黑暗中，只有他的声音）这就是长安。这就是我的——将进酒。')

add_stage_direction('第一幕结束。黑暗。')

add_para('── 第一幕结束 ──', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ── 幕间 ──
doc.add_heading('幕间：《梦游天姥》', level=2)
add_para('时长：约4分钟 | 纯器乐+舞蹈 | 李白(A)与李白(B)首次同台', bold=True)

add_stage_direction('纱幕降下——悬挂在舞台中线位置。纱幕后，月亮灯箱重新亮起——亮度约一半。')
add_stage_direction('月（舞者）第一次在纱幕后出现。她穿着简单的白色现代舞长裙——与全剧所有人的古装形成鲜明对照。她的动作是现代的——蜷缩、伸展、旋转、坠落。她不是任何具体的人——她就是"月"本身。')
add_stage_direction('音乐：键盘手切到ambient合成器音色+大量混响。古筝（国乐手）的摇指经过延迟效果器处理——制造shoegaze式的音墙。')
add_stage_direction('李白(A)从舞台左侧上——他仍然是第一幕的打扮，但动作变得缓慢，像是走在水中。他走到纱幕前，伸手——试图触碰纱幕后月的舞姿。手穿不过去。')
add_stage_direction('纱幕后，另一个人影出现。')

add_stage_direction('——李白(B)从舞台右侧深处出现。这是他全剧的第一次亮相。白色长衫已经磨损，颜色发灰。头发散开。没有剑。他看起来像是走了很远的路。')

add_stage_direction('李白(A)在纱幕前。李白(B)在纱幕后。月在他们之间舞蹈。两人隔着一层纱——对视。')

add_stage_direction('李白(A)开口——唱出《大鹏赋》的原版旋律。清亮。少年气。')
add_stage_direction('李白(B)试图跟唱——但他的嗓子已经沙哑了。Same melody, different key——他比李白(A)低了整整一个八度。他的声音不是唱出来的——像是从身体深处渗出来的。')

add_stage_direction('两个人隔纱合唱同一段旋律——少年在上，沧桑在下。月在他们之间——她触碰纱幕，纱幕被灯光打出一个手印的影子。')
add_stage_direction('月亮灯箱缓慢调暗——从半亮到微光。音乐在一个不协和的长音上悬停——然后消失。')
add_stage_direction('纱幕升。李白(A)已经不见了——只有李白(B)站在舞台中央。月退入屏风后。第一幕结束，第二幕开始。')

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 第二幕台本
# ═══════════════════════════════════════════════════
doc.add_heading('十六、第二幕台本：明月', level=1)

# ── 第八场：梁园吟 ──
doc.add_heading('第八场：《梁园吟》', level=2)
add_para('时长：约8分钟 | 演唱：李白(B)+宗倩 | 风格：布鲁斯民谣+古琴', bold=True)

add_stage_direction('灯光：清冷的蓝色面光，但比终南山的蓝更暖一些——像秋日午后的天空。屏风重新立起，两扇翻蓝色面，四扇保持白色。')
add_stage_direction('一面单独的屏风推到台前——白布面朝向观众。旁边的桌上放着毛笔和一碗清水。')
add_stage_direction('宗倩——约35岁，次女高音。她穿着素净的道袍，比玉真公主的更朴素。她不是一个"漂亮"的女人——她是一个"有重量"的女人。她的每一个动作都是从容的。')
add_stage_direction('李白(B)从舞台左侧上。距离离开长安已经过去了好几年。他的身体语言完全变了——不再是少年的轻盈，而是一种沉淀过的重量。')

add_dialogue('宗倩', '（正在看屏风——上面有字）这是你写的？')
add_dialogue('李白(B)', '（走近，看了一眼）昨天。喝多了。（自嘲地笑）老毛病。')
add_dialogue('宗倩', '"我本楚狂人，凤歌笑孔丘。"（转头看他）你不是楚人。')
add_dialogue('李白(B)', '但我是狂人。这一点——（他指了指墙上潦草的字迹）——没有骗你。')

add_stage_direction('宗倩没有笑。她认真地看了他一会儿。')

add_dialogue('宗倩', '你知道这面墙现在值多少钱吗？')
add_dialogue('李白(B)', '一面破墙？')
add_dialogue('宗倩', '因为你写在上面了。昨天有人出千金要买这面墙。')

add_stage_direction('李白(B)愣住。他看着墙上的自己写的字——像是不认识它们了。')

add_dialogue('李白(B)', '（轻轻摇头）诗——不应该变成钱。')
add_dialogue('宗倩', '那诗应该变成什么？')
add_dialogue('李白(B)', '（想了想）变成——风。吹过一个人的耳朵，然后就散了。如果她记住了——那是她的事。不是诗的事。')

add_stage_direction('宗倩拿起桌上的毛笔，蘸了清水。在李白那首诗的下方，她写下了自己的两句。清水在白布上反光——字迹清晰了几秒，然后开始慢慢变淡。')

add_dialogue('宗倩', '（边写边说）你这种人——不该一个人走路。一个人走——风会把你吹散。需要有一个人——不是拉住你，而是——跟你一起被风吹。')

add_stage_direction('她放下笔。')

add_dialogue('李白(B)', '你在跟我说——你是那个人？')
add_dialogue('宗倩', '我在跟你说——我也是修道的人。修道的人不拉人。修道的人——（指了指他写在墙上的字，又指了指自己刚写的字）——把字写在墙上。让风决定。')

add_stage_direction('唱《梁园吟》。')

add_lyric_line('【宗倩 主歌】', 0.5)
add_lyric_line('你写在墙上的字——我读了三遍。')
add_lyric_line('第一遍——我看到的是诗。')
add_lyric_line('第二遍——我看到的是写诗的人。')
add_lyric_line('第三遍——我看到的是——那个人的孤独。')

add_lyric_line('【李白(B) 回应】', 0.5)
add_lyric_line('孤独不是病——孤独是我的皮肤。')
add_lyric_line('脱不下来——也不需要脱。')
add_lyric_line('可是你说——你说——你也是修道的人。')
add_lyric_line('道是什么？')

add_lyric_line('【宗倩 副歌】', 0.5)
add_lyric_line('道——是水。水没有形状——但它能穿过所有的石头。')
add_lyric_line('你也是一块石头——你知道。')
add_lyric_line('只是没有人——像水一样——穿过你。')

add_lyric_line('【李白(B) 桥段】', 0.5)
add_lyric_line('（看着她写在墙上的字——字迹正在消失）')
add_lyric_line('你的字——也在消失。跟我的诗一样。')
add_lyric_line('墙留不住诗——但看诗的人——')

add_lyric_line('【宗倩 回应】', 0.5)
add_lyric_line('看诗的人——就是诗的归宿。')

add_lyric_line('【二人合唱 尾段】', 0.5)
add_lyric_line('（李）梁园的墙——过一百年就倒了。')
add_lyric_line('（宗）可是风——风一直在吹。')
add_lyric_line('（合）把字写在风里——风不会倒。')

add_stage_direction('歌曲结束时，屏风上的字已经完全消失。两人并肩看着空白的墙——然后相视。不是浪漫的凝视——是两个人同时意识到：他们说的是同一种语言。')
add_stage_direction('转场。')

doc.add_paragraph()

# ── 第九场：醉眠秋共被 ──
doc.add_heading('第九场：《醉眠秋共被》', level=2)
add_para('时长：约8分钟 | 演唱：李白(B)+杜甫 | 风格：吉他二重奏，兄弟情', bold=True)

add_stage_direction('两把椅子。两个酒碗。灯光是最简单的暖色面光——这一场全剧最安静。不需要任何布景效果。')
add_stage_direction('杜甫——约28岁，男中音。比李白小了十几岁。他穿着朴素的布衣。他的眼睛里有一种东西——崇拜。不是追星的崇拜，而是——一个人找到了自己的"答案"的那种崇拜。')

add_stage_direction('李白(B)已经坐在椅子上。杜甫站在旁边——不敢坐。')

add_dialogue('李白(B)', '（抬头看他）坐下。我又不是皇帝。')
add_dialogue('杜甫', '（坐下，但只坐了半张椅子）李先生……')
add_dialogue('李白(B)', '（打断）叫我太白。或者老李。或者——嘿。不要叫"先生"。我离"先生"还差——（用手指比了一下）——这么厚的一本诗集。')
add_dialogue('杜甫', '（认真地）您已经写了那么多——')
add_dialogue('李白(B)', '（把酒碗推给他）喝。喝完再说。')

add_stage_direction('杜甫接过碗，一饮而尽。呛了一下。李白(B)大笑。')

add_dialogue('杜甫', '（擦了擦嘴，忽然变得非常认真）太白——我从小就背您的诗。十五岁的时候——读到"大鹏一日同风起"。那个晚上我没睡着。')
add_dialogue('李白(B)', '（酒碗停在半空）为什么？')
add_dialogue('杜甫', '因为我想——一个人怎么可以这样写诗？像他不是在写——像是那些句子本来就在天上，他只是把它们接下来。')

add_stage_direction('李白(B)放下碗。他看了杜甫很久。')

add_dialogue('李白(B)', '（声音轻了）你知道吗——很久了。很久没有人跟我说——我的诗，对他们意味着什么。长安那群人——他们要我写诗，像点菜一样：李白，来一首牡丹。李白，来一首贵妃。我的诗——对他们来说是一个东西。就像——（指了指酒碗）——一碗酒。喝下去，暖和一下，忘了。')
add_dialogue('杜甫', '不是这样的。至少——对我来说不是。')
add_dialogue('李白(B)', '对你是什么？')
add_dialogue('杜甫', '（低头想了一会儿）是——方向。我知道我要去哪里——但我不知道怎么去。您的诗告诉我——不用知道怎么去。先出发。路会在脚下——自己出现。')

add_stage_direction('唱《梦李白》。杜甫独唱——吉他清音伴奏（吉他手切到木吉他音色）。')

add_lyric_line('【主歌】', 0.5)
add_lyric_line('死别已吞声——生别常恻恻。')
add_lyric_line('江南瘴疠地——逐客无消息。')
add_lyric_line('故人入我梦——明我长相忆。')
add_lyric_line('恐非平生魂——路远不可测。')

add_lyric_line('【副歌】', 0.5)
add_lyric_line('你是大鹏——可大鹏也会累吗？')
add_lyric_line('你飞过那么多地方——有没有一个地方——')
add_lyric_line('让你想停下来——哪怕只是一天——')
add_lyric_line('让风也歇一歇。')

add_lyric_line('【桥段——李白(B) 加入和声】', 0.5)
add_lyric_line('（杜）千秋万岁名——寂寞身后事。')
add_lyric_line('（李）身后事——我不要身后事。')
add_lyric_line('（杜）可是你不在了——诗还在。')
add_lyric_line('（李）诗在——就够了吗？')
add_lyric_line('（杜）（停顿）够了。诗在——你就在。')

add_stage_direction('歌曲结束。杜甫站起来，向李白(B)深深一拜。李白把他拉起来——给了他一个拥抱——重重的。')

add_dialogue('李白(B)', '（在杜甫耳边，低声）子美。你要写得比我好。')
add_dialogue('杜甫', '（被这个称呼震住了——"子美"只有最亲近的人才会叫）我……')
add_dialogue('李白(B)', '你一定要写得比我好。因为——（松开他，看着他的眼睛）——大鹏会老。但诗不会老。你多写一首——就多一只飞出去的鸟。')

add_stage_direction('杜甫站在原地。李白(B)转身走了。灯光缓慢收暗——只剩下杜甫的身影。')
add_stage_direction('转场。')

doc.add_paragraph()

# ── 第十场：渔阳鼙鼓 ──
doc.add_heading('第十场：《渔阳鼙鼓》', level=2)
add_para('时长：约9分钟 | 演唱：杨玉环（独唱）| 风格：工业摇滚+破碎花腔', bold=True)

add_stage_direction('黑暗中。中国大鼓——第一声。沉重而遥远。然后是第二声。第三声。越来越近，越来越响。')
add_stage_direction('鼓手独奏30秒——从缓慢到疯狂。鼓声中，群演走到屏风前，一扇一扇地把屏风推倒。每一声重鼓倒一扇。当六扇屏风全部倒地——杨玉环站在舞台中央。')
add_stage_direction('她穿着一件已经被撕裂的红色长裙——那是《沉香亭》中同一件裙子，但此刻被撕破、沾上泥土和血渍。她的头发散开。她的脚是赤的。')
add_stage_direction('灯光——深红色顶光+侧光白——两种色调同时打在她身上，让她看起来同时在被照亮和被撕开。')

add_dialogue('杨玉环', '（对着虚空——也许是对着玄宗，也许是对着自己）他们说——是我毁了盛世。（她笑了一声——那笑声不是笑声）一座帝国，毁在一个女人身上——这帝国也太轻了。')

add_stage_direction('她抬头——月亮灯箱还在。亮度调到新月模式——只剩一弯细细的光。')

add_dialogue('杨玉环', '李白临走前——他看了我一眼。我知道那一眼的意思。他在说——你也该走。可是我能走到哪里去？他是风——他可以去任何地方。我是牡丹——我离开沉香亭的土——就会死。')

add_stage_direction('唱《霓裳》——全剧最沉重的一首歌。工业摇滚的降调riff + 她破碎的花腔。')

add_lyric_line('【主歌 — 低声，几乎是在说话】', 0.5)
add_lyric_line('霓裳羽衣——穿了十五年。')
add_lyric_line('脱下来的时候——才发现——')
add_lyric_line('不是衣服重。是身体轻。')
add_lyric_line('轻到——可以被一阵风吹走。')

add_lyric_line('【副歌前 — 花腔开始，但音不准——有意地"破"】', 0.5)
add_lyric_line('他们说——倾国倾城！')
add_lyric_line('倾国——倾城——倾的是什么国？倾的是什么城？')
add_lyric_line('是一座不需要任何理由就可以把罪——')
add_lyric_line('推给一朵花的国。一座会杀死自己的牡丹的城。')

add_lyric_line('【副歌 — 工业摇滚riff进入 — 她不再试图"美"】', 0.5)
add_lyric_line('我不需要你原谅！我不需要你们记住！')
add_lyric_line('霓裳——不——是——我的！')
add_lyric_line('（"我的"两个字——嘶吼出来的花腔高音——然后陡然坠落八度）')
add_lyric_line('它是你们给我穿上的——现在我脱下来——')
add_lyric_line('还给你们。')

add_lyric_line('【桥段 — 只剩键盘和弦乐pad，她疲惫地】', 0.5)
add_lyric_line('李白——你在哪里？')
add_lyric_line('你说过云想衣裳花想容——')
add_lyric_line('可现在没有云——没有花——没有人想我——')
add_lyric_line('只有马嵬坡的风——很冷。')

add_lyric_line('【尾声 — 白绫落下（舞台上用一条白色长纱巾代替）】', 0.5)
add_lyric_line('霓裳——羽衣——')
add_lyric_line('（最后一个音——花腔的supermoney高音——然后——）')
add_lyric_line('（割断。沉默。）')

add_stage_direction('灯光——从深红骤变为全白。然后全灭。3秒完全的黑暗。')
add_stage_direction('再亮时——杨玉环已经不在。只有地上那条白色纱巾。月亮灯箱缓慢调暗至近乎全黑。')
add_stage_direction('转场。')

doc.add_paragraph()

# ── 第十一场：夜郎 ──
doc.add_heading('第十一场：《夜郎》', level=2)
add_para('时长：约10分钟 | 演唱：李白(B)+宗倩 | 风格：钢琴独白+二胡', bold=True)

add_stage_direction('全剧最暗的场景。一扇倒地的屏风斜靠在另一扇上——构成一个临时的牢笼。一束极窄的顶光从屏风缝隙漏下，照亮李白(B)半张脸。')
add_stage_direction('他坐在地上。胡须长了。白色长衫已经脏污。')

add_dialogue('李白(B)', '（像是自言自语）我这一生——做过很多错误的决定。但投永王的幕府——（苦笑）——是错得最离谱的。他跟我说——"太白先生，我们来拯救大唐。"我信了。我居然——信了。（长久的沉默）我父亲说过——你身上流的不是只有中原的血。你看得比他们远。可是有时候——看得远的人，反而看不见脚下的坑。')

add_stage_direction('宗倩从舞台左侧上。她比上一场更瘦了——来回奔波。但她仍然是从容的。')

add_dialogue('宗倩', '（站在"牢笼"外）我在外面——该找的人都找了。该花的钱都花了。还需要更多。')
add_dialogue('李白(B)', '宗倩——不要再花你的钱了。')
add_dialogue('宗倩', '我的钱——我决定怎么花。')
add_dialogue('李白(B)', '你把它花在——一个快六十岁的老头子身上——他可能死在半路上——')
add_dialogue('宗倩', '（打断）死在半路上——和死在狱中——是两回事。死在半路上——你至少还在路上。')

add_stage_direction('李白(B)抬起头，透过屏风的缝隙看着她。')

add_dialogue('李白(B)', '你为什么……')
add_dialogue('宗倩', '太白。我修道几十年——不是因为我想长生不老。是因为我在找一个答案——人活着——到底什么是值得的。（她把手伸过屏风的缝隙，触碰他的脸）我找到了。不是道——是你。不是因为你是什么天才——是因为——你不愿意装成你不是的那个人。到死都不愿意——哪怕装一下能保命。')

add_stage_direction('唱《千金赎》。宗倩独唱——键盘手切到纯钢琴音色。国乐手拉二胡。')

add_lyric_line('【主歌】', 0.5)
add_lyric_line('梁园的墙——不值千金。')
add_lyric_line('值千金的是——你写在墙上的那八个字：')
add_lyric_line('"我本楚狂人——凤歌笑孔丘。"')
add_lyric_line('那不是诗——那是你的骨头。')

add_lyric_line('现在他们要打断你的骨头——')
add_lyric_line('因为他们怕——怕一根不肯弯的骨头。')
add_lyric_line('我不怕。我卖掉所有——连道袍都可以卖——')
add_lyric_line('因为道不在衣服里。（停顿）道在我做过的事里。')

add_lyric_line('【副歌】', 0.5)
add_lyric_line('千金的墙——我买下来。千金的命——我赎回来。')
add_lyric_line('不是因为你值得——')
add_lyric_line('是因为——我选择你值得。')
add_lyric_line('选择——就是我修了一辈子的道。')

add_lyric_line('【副歌 重复——二胡旋律与钢琴对话】', 0.5)
add_lyric_line('如果他们说——你罪不可赦——')
add_lyric_line('那我陪你一起罪。')
add_lyric_line('如果他们说——你老了就认命吧——')
add_lyric_line('（钢琴渐弱，二胡独奏一句——然后）')
add_lyric_line('那我说——老了的狂人——还是狂人。')

add_stage_direction('歌曲结束。宗倩把手从屏风缝隙中收回。转身——走向舞台后方。李白(B)的顶光收到最小——只剩一只眼睛的光。')
add_stage_direction('转场。')

doc.add_paragraph()

# ── 第十二场：轻舟已过 ──
doc.add_heading('第十二场：《轻舟已过》', level=2)
add_para('时长：约7分钟 | 演唱：李白(B) | 风格：明亮的公路摇滚', bold=True)

add_stage_direction('与上一场形成最强烈的反差。灯光——开放白面光+暖色顶光。全剧最明亮的段落之一。')
add_stage_direction('一面倒地屏风平放在舞台上——当作"船"。群演8人在两侧，手持蓝色长布上下飘动——长江的波浪。节奏是轻快的。')
add_stage_direction('李白(B)站在"船"上。他老了。脏了。但他的眼睛——重新亮起来了。')

add_dialogue('李白(B)', '（对着观众，语气像一个刚从恶作剧中脱身的孩子）他们不要我了。皇帝说——算了，让他走吧。一个快六十岁的老头子，流放到夜郎——走了一年还没走到——他还能造反？（大笑）于是我就——回来了！')

add_stage_direction('他张开手臂，像要拥抱整条长江。')

add_dialogue('李白(B)', '我从白帝城上船的时候——天还没亮。江面上全是雾。船夫问我去哪里——我说——（忽然顿住了）我忽然不知道该怎么说。因为我第一次发现——我不需要去任何地方了。我只需要——在路上。')

add_stage_direction('唱《轻舟已过》。')

add_lyric_line('【主歌 — 轻快的摇滚节奏】', 0.5)
add_lyric_line('朝辞白帝彩云间——千里江陵一日还。')
add_lyric_line('两岸猿声啼不住——轻舟已过万重山。')

add_lyric_line('【扩展段 — 现代翻译式扩展】', 0.5)
add_lyric_line('早晨的白帝城——还在云里。')
add_lyric_line('可是我已经在千里之外——')
add_lyric_line('猴子还在岸上叫唤——')
add_lyric_line('我的船已经过了——一万座山！')

add_lyric_line('【副歌】', 0.5)
add_lyric_line('我不是逃——我只是不需要留在那里。')
add_lyric_line('他们关不住我——因为关住一个人——')
add_lyric_line('需要那个人先同意被关。')
add_lyric_line('而我——（笑）——从来没有同意过！')

add_lyric_line('轻舟——轻舟——轻舟已过万重山——！')
add_lyric_line('所有压在我身上的山——')
add_lyric_line('从后面看——只是风景。')

add_lyric_line('【尾声】', 0.5)
add_lyric_line('水流向海——船跟着水——我跟船——')
add_lyric_line('不需要桨——不需要帆——')
add_lyric_line('只要还在漂——我就还没有输。')

add_stage_direction('歌曲结束——李白(B)从"船"上跳下来。蓝布落在地上——变成了静止的水。他回望了一眼。然后继续往前走。')
add_stage_direction('灯光缓慢转暗——从明亮的白天渐变为黄昏。转场。')

doc.add_paragraph()

# ── 第十三场：月下独酌 ──
doc.add_heading('第十三场：《月下独酌》', level=2)
add_para('时长：约8分钟 | 演唱：李白(B)+月（舞）| 风格：古筝+电吉他二重奏', bold=True)

add_stage_direction('全剧最安静也最孤独的一场。灯光——月光蓝面光+月亮灯箱调至满月。没有屏风——它们全部被推到舞台最后方，杂乱地堆在一起。')
add_stage_direction('舞台中央——一把椅子，一个小酒壶放在地上。李白(B)——他已经很老了。白色长衫多处打补丁。头发全白。他坐在椅子上，不是在休息——是在等。')
add_stage_direction('良久。他弯腰拿起酒壶——空的。倒了倒——最后一滴落在舞台上。')

add_dialogue('李白(B)', '（非常轻地）没有了。连酒——都没有了。')

add_stage_direction('他抬头看月亮灯箱。月亮是满的——明亮的。')

add_dialogue('李白(B)', '（对月亮说话）你还在。你一直都在。从我在四川——（指了指自己的膝盖，像在指一个小孩的高度）——这么高的时候——抬头看，是你。在金陵——在长安——在夜郎的狱中——抬头看——还是你。所有人都不在了。父亲。许宛。杨玉环。宗倩——（他顿了顿，声音更轻）她卖掉了最后一件值钱的东西。她跟我说不要紧。她说——道不是东西。道是我做过的事。（停顿）她说这话的时候——还在笑。')

add_dialogue('李白(B)', '（对月亮）你知道吗——我写了一千首诗。但有一首——我一直写不出来。就是给你写的。因为每次我试着写——都发现不是我在写月亮。是月亮在写我。')

add_stage_direction('月（舞者）从月亮灯箱下方缓步走出。她没有穿鞋。她的白色现代舞长裙在月光下几乎是透明的。她走向李白(B)——每一步都很慢，像在水里。')

add_stage_direction('李白(B)站起来。和她面对面。')

add_dialogue('李白(B)', '我认识你。从我五岁的时候——你就一直在。但我从来没有——像现在离你这么近。')

add_stage_direction('他向月伸出手。月也伸出手——但她的手停在半空，离他的手只差一厘米。不触碰。')

add_stage_direction('唱《月下独酌》——古筝（国乐手）+ 电吉他清音二重奏。全剧唯一一首没有鼓的歌。')

add_lyric_line('【主歌】', 0.5)
add_lyric_line('花间一壶酒——独酌无相亲。')
add_lyric_line('举杯邀明月——对影成三人。')
add_lyric_line('月既不解饮——影徒随我身。')
add_lyric_line('暂伴月将影——行乐须及春。')

add_lyric_line('你从来不说话——可你什么都说了。')
add_lyric_line('你不喝我的酒——可你一直陪着我。')
add_lyric_line('从我五岁——到六十岁——')
add_lyric_line('（停顿——古筝与吉他对话式的交替独奏4小节）')
add_lyric_line('只有你——从来没变过。')

add_lyric_line('【副歌 — 他对着月唱——月绕着他舞蹈】', 0.5)
add_lyric_line('他们都走了——父亲，许宛，杨玉环——')
add_lyric_line('宗倩说她"选择我值得"——可是值得吗？')
add_lyric_line('我不知道。我只知道——我一直在飞——')
add_lyric_line('现在，飞不动了。（他看着月亮）可是我看到——')
add_lyric_line('你在水里。你在水里面。你在叫我——')

add_lyric_line('【尾声 — 他走向台口。观众席方向——那是"水"的方向】', 0.5)
add_lyric_line('我醉欲眠——卿且去——')
add_lyric_line('（他停了很久。然后——像是对自己，也像是对所有人）')
add_lyric_line('大鹏——飞了一辈子——')
add_lyric_line('最后才发现——海不在远方——')
add_lyric_line('海在——（他低头看着自己的手心——仿佛那里有一片海）——')
add_lyric_line('海在这里面。一直都是。')

add_stage_direction('顶光——只照亮李白(B)仰面向上的脸。他伸出手——这一次，月握住了他的手。两人慢慢走向台口。')
add_stage_direction('然后——顶光熄灭。完全的黑暗。')
add_stage_direction('3秒.')
add_stage_direction('月亮灯箱缓缓亮起——满月。但舞台上李白(B)已经不在了。只有月独自站在舞台中央——她的独舞，缓慢的、庄严的，像一场看不见的送别。')
add_stage_direction('月最后走到月亮灯箱下方。她抬起手臂——灯箱的亮度慢慢降到最低。然后她也消失了。')

doc.add_paragraph()

# ── 终曲：骑鲸 ──
doc.add_heading('终曲：《骑鲸》', level=2)
add_para('时长：约5分钟 | 演唱：全员 | 风格：管弦乐合唱+主题重现', bold=True)

add_stage_direction('黑暗中——海浪声。不是真实的录音——是键盘手的合成器pad+合唱团的低声哼鸣。')
add_stage_direction('月亮灯箱缓慢亮起——不是满月，而是一种柔和的、均匀的亮光，像清晨的海面。')
add_stage_direction('六扇屏风重新立起来了——但每一扇颜色不同：白、金、蓝、红、绿、黄——像李白一生走过的所有颜色。')

add_stage_direction('第一个登场的是李客。他从白色屏风后走出。然后是阿素——从绿色屏风后。玉真公主——从金色屏风后。许宛——从黄色屏风后。杜甫——从蓝色屏风后。杨玉环——从红色屏风后。宗倩——从白色屏风后。')
add_stage_direction('他们每人站定一个位置。不看向彼此——所有人都看着舞台中央的空地。那是李白应该站的地方。')

add_stage_direction('然后——从舞台后方，李白(A)出现。白色长衫整洁，佩剑，20岁的模样。他从屏风之间穿过，走到中央。然后他停下——转身——伸出手。')
add_stage_direction('李白(B)从同一方向出现——灰色长衫，散发，无剑。他走向李白(A)——握住那只手。两人并肩站在舞台中央。')

add_lyric_line('【合唱 — 管弦乐+合唱团，速度 ♩=60】', 0.5)
add_lyric_line('（全团轻声）大鹏——一日同风起——')
add_lyric_line('（全团渐强）扶摇直上——九万里——')

add_lyric_line('【李白(A) 独唱 — 少年旋律，上方声部】', 0.5)
add_lyric_line('我是风——我是飞过蜀道的风——')
add_lyric_line('我是吹过金陵酒肆的风——')
add_lyric_line('我吹过终南山——吹过安陆的桃花——')
add_lyric_line('我吹过长安——吹过沉香亭的牡丹——')

add_lyric_line('【李白(B) 独唱 — 沧桑旋律，下方声部，与(A)形成八度和声】', 0.5)
add_lyric_line('我是风——我是被夜郎的石头挡住了的风——')
add_lyric_line('我是吹过长江——吹过万重山的风——')
add_lyric_line('我是吹了一辈子——终于——')
add_lyric_line('（他看了一眼月亮灯箱）——被月亮接住了的风。')

add_lyric_line('【全体合唱 — 两个李白+所有角色+合唱团】', 0.5)
add_lyric_line('大鹏一日同风起——扶摇直上九万里！')
add_lyric_line('假令风歇时下来——犹能簸却沧溟水！')
add_lyric_line('他不是坠落——他骑上了鲸鱼！')
add_lyric_line('大鹏回到了风里——回到了海上——')
add_lyric_line('回到了每一首诗的第一个字里！')

add_stage_direction('合唱推向最强音——然后——骤停。')
add_stage_direction('沉默。两拍。所有演员转身，走向舞台后方。两个李白走在最后——他们的背影在月亮灯箱的光芒中越来越小。')
add_stage_direction('月亮灯箱——缓缓熄至全黑。剧场全暗。')
add_stage_direction('5秒。')

doc.add_paragraph()

# ── 谢幕 ──
doc.add_heading('谢幕：《诗活着》', level=2)
add_para('时长：约4分钟 | 演唱：全员 | 风格：流行摇滚', bold=True)

add_stage_direction('灯光重新亮起——这次是明亮的暖色全光。所有演员从屏风后走出来——但他们已经换掉了戏服的外袍，露出里面的现代便装。T恤、卫衣、牛仔裤。')
add_stage_direction('每个人手里拿着手机——屏幕亮着。')
add_stage_direction('李白(A)走到台口——他现在穿着普通的黑色T恤和牛仔裤。但腰间——还挂着那把剑。')

add_dialogue('李白(A)', '（对观众，不再扮演"李白"，而是一个演完了李白的学生）他的诗——一共流传下来——大约有一千首。也就是说——一千三百年来——每一天——这个世界上的某个角落——都有人在读他的诗。')

add_dialogue('李白(A)', '（举起手机——屏幕上是"床前明月光"）今晚——轮到你们了。')

add_stage_direction('唱《诗活着》。')

add_lyric_line('【全体合唱】', 0.5)
add_lyric_line('他走了——可是月亮还在照。')
add_lyric_line('他死了——可是诗没有死。')
add_lyric_line('长安的宫殿——一百年前就倒了。')
add_lyric_line('可是"床前明月光"——昨晚还在被一个孩子念出来。')

add_lyric_line('他把字写在墙上——墙倒了。')
add_lyric_line('他把字写在水里——水干了。')
add_lyric_line('他把字写在风里——（笑）——风还在吹。')
add_lyric_line('吹了一千三百年——吹到了今晚——吹到了你的耳朵里。')

add_lyric_line('【副歌】')
add_lyric_line('诗活着——在每一个背诗的小孩嘴里！')
add_lyric_line('诗活着——在每一个想家的夜晚！')
add_lyric_line('诗活着——在你说不出来的话里！')
add_lyric_line('当他写下来——世界就多了一个出口。')

add_lyric_line('【尾声 — 所有人举起手机，屏幕对着观众——屏幕上滚动着李白诗句】', 0.5)
add_lyric_line('诗活着。你也是。')

add_stage_direction('所有演员鞠躬。月亮灯箱在舞台后方独自亮着——那是全剧最后一个画面。')

add_para('── 全剧终 ──', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 第四部分：样曲
# ═══════════════════════════════════════════════════
add_para('第四部分：样曲', 'Normal', bold=True, font_size=16, font_name='微软雅黑',
         color=RGBColor(0x1a, 0x1a, 0x2e), alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

doc.add_heading('十七、《将进酒》——完整歌词与编曲说明', level=1)

add_para('选这首歌作为样曲的理由：它是全剧第一幕的高潮，是李白"摇滚精神"最集中的爆发点，也是原诗知名度最高的作品。原诗天然具备verse-chorus-bridge结构，改编为摇滚歌曲的空间非常充分。', italic=True)

doc.add_heading('17.1 基本信息', level=2)
add_table_with_data(
    ['项目', '内容'],
    [
        ['曲名', '《将进酒》'],
        ['位置', '第一幕第七场（第一幕高潮）'],
        ['演唱', '李白(A) + 合唱团 + 全体演员'],
        ['风格', '硬摇滚 + 中国大鼓 + 琵琶竞奏'],
        ['速度', '♩=132（终副歌加速至 ♩=144）'],
        ['调性', 'E小调（适合摇滚吉他开放弦共振）'],
        ['时长', '约 5\'00"（含前奏尾奏）'],
    ]
)

doc.add_heading('17.2 歌曲结构', level=2)
add_table_with_data(
    ['段落', '小节数', '内容', '配器'],
    [
        ['引子', '4小节', '中国大鼓三声重击 + 电吉他长音Feedback', '大鼓 → 电吉他失真'],
        ['主歌 A1', '8小节', '李白独唱"君不见黄河之水天上来"', '贝斯+键盘轻铺+琵琶单音'],
        ['主歌 A2', '8小节', '声压渐强"人生得意须尽欢"', '架子鼓进入+吉他清音分解'],
        ['合唱过渡', '4小节', '合唱团"三百杯×3"', '全编制渐强→骤停'],
        ['桥段（吟诵）', '8小节', '"岑夫子丹丘生"近似吟诵', '鼓走trap groove+琵琶点缀'],
        ['副歌', '16小节', '大合唱"钟鼓馔玉不足贵"', '全编制+合唱团'],
        ['间奏', '16小节', '电吉他solo + 琵琶同度齐奏', '全编制（标志性"古今碰撞"）'],
        ['独唱段', '8小节', '"五花马千金裘"', '贝斯+鼓，最简编配'],
        ['终副歌', '16+4小节', '加速重复副歌→结尾三和弦', '全编制+全员+速度推到144'],
    ]
)

doc.add_heading('17.3 完整歌词', level=2)

add_para('说明：原诗部分以标记【原诗】，扩展创作部分标记为【新词】。作曲时，原诗段落走五声调式旋律，新词段落融入布鲁斯♭3♭7音。', italic=True)

doc.add_heading('【引子】', level=3)
add_stage_direction('电吉他长音Feedback从低声部爬升 + 中国大鼓三声重击 + 琵琶扫弦引入')
add_lyric_line('（器乐 4小节）')

doc.add_heading('【主歌 A1 — 独唱，乐队轻伴奏】', level=3)
add_lyric_line('【原诗·李白】')
add_lyric_line('君不见——黄河之水天上来——')
add_lyric_line('奔流到海不复回。')
add_lyric_line('君不见——高堂明镜悲白发——')
add_lyric_line('朝如青丝暮成雪。')
add_stage_direction('——吉他失真进入，功率和弦。鼓手从中国大鼓回到架子鼓——')

doc.add_heading('【主歌 A2 — 声压渐强】', level=3)
add_lyric_line('【原诗·李白】')
add_lyric_line('人生得意须尽欢——莫使金樽空对月！')
add_lyric_line('天生我材必有用——千金散尽还复来！')
add_lyric_line('烹羊宰牛且为乐——会须一饮三百杯！')
add_stage_direction('——合唱团加入，叠在主旋律上方——')

doc.add_heading('【合唱过渡】', level=3)
add_lyric_line('【新词·发展】')
add_lyric_line('（合唱团，渐强）三百杯！三百杯！三——百——杯！')
add_stage_direction('——乐队骤停。只留踩镲嘶嘶声。灯光压暗至追光只打李白。——')

doc.add_heading('【桥段 — 近似吟诵，鼓走Trap groove】', level=3)
add_lyric_line('【原诗·李白】')
add_lyric_line('岑夫子——丹丘生——将进酒——杯莫停。')
add_lyric_line('与君歌一曲——请君为我倾耳听。')
add_lyric_line('【新词·延伸】')
add_lyric_line('（李白，对观众，近似说话）')
add_lyric_line('你们听好了——这首诗——不是写给你们的。')
add_lyric_line('是写给一千年以后——那个也在半夜里——')
add_lyric_line('对着月亮发呆的人。那个人——是不是你？')
add_stage_direction('——全编制回归。金色+红色顶光同时亮。全体演员站到台前。——')

doc.add_heading('【副歌 — 大合唱】', level=3)
add_lyric_line('【原诗·李白 —— 李白(A)主唱，全员和声】')
add_lyric_line('钟鼓馔玉不足贵——但愿长醉不愿醒！')
add_lyric_line('古来圣贤皆寂寞——惟有饮者留其名！')
add_lyric_line('陈王昔时宴平乐——斗酒十千恣欢谑！')
add_lyric_line('主人何为言少钱——径须沽取对君酌！')
add_lyric_line('【新词·副歌扩充】')
add_lyric_line('（全团，在"对君酌"之后不停，直接推上去——）')
add_lyric_line('喝——！不是为了忘记！是为了记住——！')
add_lyric_line('记住今晚——记住这个月亮——记住还在路上的自己——！')

doc.add_heading('【间奏 — 电吉他solo + 琵琶同度齐奏 16小节】', level=3)
add_stage_direction('全剧标志性的"古今碰撞"声音。电吉他失真音色与琵琶扫拂走同一条旋律线——一个撕裂、一个清脆。节奏组走复合节奏——中国大鼓与架子鼓交替，切分错位。')
add_stage_direction('灯光——金色与红色交替闪烁，像着了火的宫殿。')

doc.add_heading('【独唱段 — 编配减到最小】', level=3)
add_lyric_line('【原诗·李白（稍作改写以适配节奏）】')
add_lyric_line('（只剩贝斯+鼓+琵琶单音）')
add_lyric_line('五花马——千金裘——')
add_lyric_line('呼儿将出换美酒——')
add_lyric_line('与尔同销——')
add_lyric_line('（鼓——重击）万——！（鼓）古——！（鼓）愁——！！！')
add_stage_direction('"万——古——愁"三字每字之间隔一拍。每一下全乐队炸一个和弦。第三声"愁"上李白(A)推至极限高音并保持——吉他Feedback长啸。')

doc.add_heading('【终副歌 — 加速至 ♩=144】', level=3)
add_lyric_line('【新词+原诗交替 —— 全团,速度推到极限】')
add_lyric_line('将进酒——杯莫停——！')
add_lyric_line('与君歌一曲——请君为我倾耳听——！')
add_lyric_line('（全团）钟鼓馔玉不足贵——但愿长醉不愿醒——！')
add_lyric_line('（全团）古来圣贤皆寂寞——惟有饮者留其名——！')
add_lyric_line('（李白(A)嘶吼）将————进————酒————！')
add_lyric_line('（全团回应）杯！莫！停！')
add_lyric_line('（李白(A)最后一次高音）将——进——酒——！！！')
add_stage_direction('——乐队最后三个强力和弦。全体演员定格。手臂高举。嘴张着。被凝固在呐喊的顶点。——')
add_stage_direction('——灯光骤灭。只有月亮灯箱亮着。——')
add_lyric_line('（李白(A)在黑暗中，轻声）这就是长安。这就是我的——将进酒。')

doc.add_heading('17.4 编曲要点（给乐队指挥看）', level=2)

add_table_with_data(
    ['乐器', '关键指令'],
    [
        ['电吉他', '主歌A1清音分解→A2功率和弦切入→副歌高增益失真→间奏solo与琵琶同度齐奏。需要Line6 Helix或同级别综合效果器，一个patch内完成音色切换。'],
        ['琵琶', '主歌A1单音点缀→A2轮指扫拂（与吉他失真形成音色对比）→间奏solo与吉他同度齐奏→终副歌狂飙扫弦。这是全剧最重要的一场琵琶——如果国乐手只能弹好一首歌，就是这首。'],
        ['中国大鼓', '引子三声重击→间奏与架子鼓交替节奏→万古愁三声重击——每一下都是剧情的标点符号。鼓手需在架子鼓和中国大鼓之间快速切换（中国大鼓放在舞台后区，架子鼓在右前乐池）。'],
        ['键盘手', '引子弦乐pad铺垫→主歌A1钢琴轻触→副歌叠加弦乐厚度→吟诵段不弹（让鼓和琵琶对话）→终副歌管弦乐全开。'],
        ['贝斯', '主歌A1/A2走简洁的根音线→副歌跟吉他走强力移动→吟诵段走trap风格的sub-bass→终副歌八分音符连续picking。'],
        ['合唱团', '从观众席两侧走上来——边唱边走。排练时需练习"边走边唱"的稳定性。副歌和声分三个声部：女高（上方三度）、女中（同度）、男声（下方八度——提前为终曲《骑鲸》的李白(B)声部做听觉伏笔）。'],
    ]
)

doc.add_heading('17.5 舞台与灯光提示', level=2)

add_table_with_data(
    ['时刻', '灯光', '动作'],
    [
        ['引子', '全暗→大鼓第一声时月亮灯箱闪一下', '李白(A)从舞台后方走向台口'],
        ['主歌A1', '追光只打李白上半身', '李白面向观众独唱'],
        ['A1→A2转换', '追光缓慢扩大', '合唱团从观众席走上来'],
        ['"三百杯"', '金色顶光渐亮', '鼓手从架子鼓走向中国大鼓'],
        ['吟诵段', '全灯压暗只剩追光', '李白走到台口最前端，Trap groove'],
        ['副歌', '金色+红色顶光全开', '全团站到台前——打破第四面墙'],
        ['间奏', '金红交替闪烁', '电吉他+琵琶走到舞台中央竞奏'],
        ['"万古愁"', '每声鼓闪一次全白光', '全演员定格三次'],
        ['终副歌', '所有灯光推到最大', '速度144，全团嘶吼'],
        ['结尾', '骤灭→只留月亮灯箱', '3秒静默→李白轻声收尾'],
    ]
)

doc.add_heading('17.6 声乐提示（给李白(A)演员）', level=2)

add_para('• 整首歌不要"美声"——这是摇滚，不是歌剧。用你说话的声音唱歌。高音部分可以破，可以撕，可以到极限——破掉比完美的假声更动人。')
add_para('• 主歌A1——用叙述的语气，像在跟一个朋友讲一个很久以前的故事。"君不见黄河之水天上来"——你不是在朗诵，你是在回忆。')
add_para('• 吟诵段——说唱式的节奏感。底鼓和贝斯给你trap的groove，你踩着那个节奏说——不用唱准音高，但要踩准每一个字的重音。')
add_para('• "万——古——愁"三个字——这是你全剧声乐的最高点，也是最危险的点。提前两场戏开始节省嗓子。唱"愁"的时候——不要把力气全用在音量上。用气息推到高音然后让它自己悬在那里——让麦克风和混响帮你。')
add_para('• 最后一个"将进酒"的嘶吼——之后你要保持口型但不出声，直到灯光灭。这是全剧第一幕的最后一个画面——不要让任何声音破坏那个留白。')

doc.add_paragraph()
doc.add_paragraph()

add_para('—— 第四部分结束。以下为附录。——', 'Normal', font_size=10, color=RGBColor(0x99,0x99,0x99), alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ── 附录 ──
doc.add_heading('附录A：原创歌词与原诗对照表', level=1)

add_table_with_data(
    ['歌曲', '原诗引用', '新词比例', '备注'],
    [
        ['《大鹏赋》', '李白《上李邕》全诗四句', '约60%为新词', '主题旋律，三次不同编曲'],
        ['《蜀道》', '化用《蜀道难》意象', '约90%新词', '原作太长，提取"难"的核心意象'],
        ['《金陵酒肆》', '化用《金陵酒肆留别》首句', '约85%新词', '扩展为男女对唱'],
        ['《终南山》', '无直接引用', '100%新词', '原创二重唱'],
        ['《安陆》', '无直接引用', '100%新词', '许宛独唱，全剧唯一完全原创'],
        ['《长安·上》', '无直接引用', '100%新词', '群像曲'],
        ['《清平调》', '李白《清平调》三首完整引用', '约30%为新词', '杨玉环花腔回应段为新词'],
        ['《脱靴》', '无直接引用', '100%新词', '短小讽刺摇滚'],
        ['《将进酒》', '李白《将进酒》全诗引用', '约40%为新词', '扩展副歌+吟诵段'],
        ['《梁园吟》', '化用"我本楚狂人"二句', '约85%新词', '宗倩与李白的问答'],
        ['《梦李白》', '杜甫《梦李白二首》节选', '约50%新词', '杜甫独唱+李白和声'],
        ['《霓裳》', "化用'霓裳羽衣'意象 | 100%新词", '杨玉环绝命歌'],
        ['《千金赎》', '无直接引用', '100%新词', '宗倩独唱'],
        ['《我本楚狂人》', '化用《庐山谣》首句', '约80%新词', '狱中Blues'],
        ['《轻舟已过》', '李白《早发白帝城》全诗', '约60%新词', '公路摇滚扩展'],
        ['《月下独酌》', '李白《月下独酌》全诗', '约50%新词', '吉他+古筝二重奏'],
        ['《骑鲸》', '化用《大鹏赋》主题', '约50%新词', '主题重现'],
        ['《诗活着》', '无直接引用', '100%新词', '谢幕曲'],
    ]
)

doc.add_heading('附录B：乐器切换时间表（给国乐手）', level=1)
add_table_with_data(
    ['场次', '乐器', '备注'],
    [
        ['序曲《大鹏赋》', '笛子（前奏引子）', '仅前奏8小节'],
        ['1《蜀道》', '笛子', '全曲'],
        ['2《金陵酒肆》', '琵琶', '全场最吃重的琵琶曲目'],
        ['3《终南山》', '古筝', '空灵琶音——允许即兴'],
        ['4《安陆》', '二胡', '温暖悲而不伤'],
        ['5《长安·上》', '琵琶', 'Funk切分节奏'],
        ['6《清平调》', '琵琶', '宫廷华丽风格'],
        ['7《将进酒》', '琵琶→中国鼓辅助', '全剧最重要的一场琵琶'],
        ['幕间', '古筝（摇指+延迟效果）', 'Shoegaze音墙'],
        ['8《梁园吟》', '古筝/古琴', '安静对话式'],
        ['9《梦李白》', '（休息）', '只有吉他，国乐手此场可休息'],
        ['10《霓裳》', '（休息）', '工业摇滚为主'],
        ['11《千金赎》', '二胡', '催泪——全剧最动人的二胡段落'],
        ['12《轻舟已过》', '笛子', '轻快明亮的公路摇滚'],
        ['13《月下独酌》', '古筝', '与吉他清音的二重奏——整部剧的音乐灵魂'],
        ['17《骑鲸》', '琵琶（轻）', '主题重现'],
        ['18《诗活着》', '琵琶', '流行摇滚'],
    ]
)

doc.add_paragraph()
add_para('—— 全文完 ——', 'Normal', font_size=14, font_name='楷体',
         color=RGBColor(0x44, 0x72, 0xC4), alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para(f'文档生成日期：{datetime.date.today().strftime("%Y年%m月%d日")}', 'Normal',
         font_size=9, color=RGBColor(0x99, 0x99, 0x99), alignment=WD_ALIGN_PARAGRAPH.CENTER)

# ── 保存 ──
output_path = r'D:\辰辰\first CC\谪仙人_李白历史摇滚音乐剧_完整策划方案.docx'
doc.save(output_path)
print(f'文档已保存至：{output_path}')
print('Done!')
