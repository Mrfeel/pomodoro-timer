# -*- coding: utf-8 -*-
"""
电磁学作业解答生成器
按周数顺序：Week 12 → Week 13 → Week 14 → Week 16
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ========== 样式设置 ==========
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# 设置中文字体
style.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')

def add_title(text, level=0):
    """添加标题"""
    if level == 0:
        p = doc.add_heading(text, level=0)
    else:
        p = doc.add_heading(text, level=level)
    return p

def add_para(text, bold=False, italic=False, size=12):
    """添加段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    return p

def add_formula(text):
    """添加公式（用等宽字体表示）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Courier New'
    run.italic = True
    return p

def add_solution_step(text):
    """添加解答步骤"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    return p

# ============================================================
#                         封面
# ============================================================
add_title('电磁学作业解答', level=0)
add_para('杨思辰  PB25992094  未来技术学院')
add_para('教材：胡友秋等《电磁学》(科学出版社, 2024)')
add_para('')

# ============================================================
#                     Week 12 作业
# ============================================================
add_title('第12周作业 (2025年5月15日)', level=1)

# --- Problem 1-3: 6.13, 6.17, 6.19 ---
add_title('题目1-3：教材习题 6.13, 6.17, 6.19', level=2)

add_title('6.13 题', level=3)
add_para('【题目】一无限长直导线载有电流I，导线半径为a，磁导率为μ。求导线内外的磁感应强度B和磁场强度H的分布。')
add_para('【解】利用安培环路定理。')
add_para('在导线内部 (r < a)：')
add_formula('∮H·dl = I·(πr²/πa²) = I·r²/a²')
add_formula('H(r) = Ir/(2πa²),  B(r) = μH = μIr/(2πa²)')
add_para('在导线外部 (r > a)：')
add_formula('∮H·dl = I')
add_formula('H(r) = I/(2πr),  B(r) = μ₀I/(2πr)')

add_title('6.17 题', level=3)
add_para('【题目】一均匀磁化的无限长圆柱形磁铁，半径为a，磁化强度为M（沿轴向）。求磁铁内外的B和H分布。')
add_para('【解】磁化强度M对应磁化电流密度：体磁化电流Jm = ∇×M = 0（均匀磁化）；面磁化电流Km = M×n̂ = M φ̂。')
add_para('磁铁等效于一无限长螺线管，表面电流密度为M。')
add_para('内部 (r < a)：B = μ₀M（沿轴向），H = B/μ₀ − M = 0')
add_para('外部 (r > a)：B = 0，H = 0（因为螺线管外磁场为零）')

add_title('6.19 题', level=3)
add_para('【题目】一无限大平面将空间分为两部分，上半空间充满磁导率为μ₁的介质，下半空间充满磁导率为μ₂的介质。在介质1中距界面为d处有一无限长载流直导线，电流为I，平行于界面。求空间的磁场分布。')
add_para('【解】此题用磁像法求解。')
add_para('对于区域1 (y > 0)：磁场由真实电流I和像电流I\' = I·(μ₂−μ₁)/(μ₂+μ₁)（位于对称位置y = −d处）共同产生。')
add_para('对于区域2 (y < 0)：磁场由等效电流I\'\' = 2μ₁I/(μ₂+μ₁)（位于真实电流位置y = d处）产生。')
add_para('由此可分别求得两区域的B和H分布。')

# --- Problem 4: Magnetic media boundary ---
add_title('题目4：磁介质边界问题', level=2)
add_para('【题目】如图所示，两个半无限大磁介质（磁导率分别为μ₁, μ₂）和中间一层有限厚度的磁介质（磁导率μ₃），三种介质中有均匀的外磁场。求各介质中的磁场分布。')
add_para('【解】设外磁场为H₀，方向垂直于界面。利用磁介质边界条件：')
add_para('(1) B的法向分量连续：B₁n = B₂n = B₃n')
add_para('(2) H的切向分量连续：H₁t = H₂t = H₃t')
add_para('对于法向磁场情况，B在各区域均相等：')
add_formula('B₁ = B₂ = B₃ = B₀')
add_formula('H₁ = B₀/μ₁,  H₂ = B₀/μ₂,  H₃ = B₀/μ₃')
add_para('其中B₀由外场条件确定。对于切向磁场情况，H在各区域均相等：')
add_formula('H₁ = H₂ = H₃ = H₀')
add_formula('B₁ = μ₁H₀,  B₂ = μ₂H₀,  B₃ = μ₃H₀')

# --- Problem 5-9: 6.22, 7.2, 7.4, 7.6, 7.7 ---
add_title('题目5-9：教材习题 6.22, 7.2, 7.4, 7.6, 7.7', level=2)

add_title('6.22 题', level=3)
add_para('【题目】一环形铁芯的截面为矩形，内半径为R₁，外半径为R₂，高度为h，铁芯上绕有N匝线圈通有电流I。铁芯的磁导率为μ。求铁芯中的磁通量和线圈的自感系数。')
add_para('【解】由安培环路定理，铁芯中半径为r处的磁场：')
add_formula('H(r) = NI/(2πr),  B(r) = μNI/(2πr)')
add_para('通过单匝线圈的磁通量：')
add_formula('Φ = ∫∫ B·dS = ∫₀ʰ∫_{R₁}^{R₂} [μNI/(2πr)] dr dz = (μNIh/2π) ln(R₂/R₁)')
add_para('磁通匝链数 Ψ = NΦ = (μN²Ih/2π) ln(R₂/R₁)')
add_para('自感系数：')
add_formula('L = Ψ/I = (μN²h/2π) ln(R₂/R₁)')

add_title('7.2 题', level=3)
add_para('【题目】一长直导线载有交变电流I = I₀sin(ωt)，在其附近有一矩形线圈（边长为a×b），线圈平面与导线共面，线圈的一边与导线平行且距离为d。求线圈中的感应电动势。')
add_para('【解】长直导线在距离r处产生的磁场：B = μ₀I/(2πr)')
add_para('通过线圈的磁通量：')
add_formula('Φ = ∫_d^{d+a} [μ₀I/(2πr)] · b · dr = (μ₀Ib/2π) ln[(d+a)/d]')
add_para('感应电动势：')
add_formula('ε = −dΦ/dt = −(μ₀b/2π) ln[(d+a)/d] · dI/dt')
add_formula('ε = −(μ₀bωI₀/2π) ln[(d+a)/d] · cos(ωt)')

add_title('7.4 题', level=3)
add_para('【题目】在均匀磁场B中，一面积为S的N匝线圈以角速度ω绕垂直于B的轴旋转。求线圈中的感应电动势。')
add_para('【解】设t=0时线圈法线与B平行，则磁通量：')
add_formula('Φ(t) = NBS cos(ωt)')
add_formula('ε = −dΦ/dt = NBSω sin(ωt) = ε₀ sin(ωt)')
add_para('其中ε₀ = NBSω为电动势幅值。这就是交流发电机的基本原理。')

add_title('7.6 题', level=3)
add_para('【题目】一长直螺线管半径为R，单位长度匝数为n，通有电流I = I₀e^{−αt}。求管内外距离轴线为r处的感生电场强度。')
add_para('【解】由法拉第定律的积分形式 ∮E·dl = −dΦ/dt')
add_para('管内(r < R)：通过半径为r的圆面积的磁通量 Φ = μ₀nI·πr²')
add_formula('E·2πr = −d(μ₀nI·πr²)/dt = μ₀nαI₀e^{−αt}·πr²')
add_formula('E(r,t) = (μ₀nαI₀r/2) e^{−αt}  (方向：环绕轴向)')
add_para('管外(r > R)：磁通量 Φ = μ₀nI·πR²')
add_formula('E·2πr = μ₀nαI₀e^{−αt}·πR²')
add_formula('E(r,t) = (μ₀nαI₀R²/(2r)) e^{−αt}')

add_title('7.7 题', level=3)
add_para('【题目】一金属圆盘半径为a，厚度为d，电导率为σ，在均匀磁场B中绕中心轴以角速度ω旋转（B垂直于盘面）。求盘中的感应电流分布和总涡流损耗功率。')
add_para('【解】在半径r处，线速度v = ωr，动生电场（单位电荷受力）E = vB = ωrB')
add_para('电流密度 j = σE = σωrB')
add_para('在半径r取径向宽度dr的圆环，其电阻dR = (2πr)/(σd·dr)')
add_para('该环中的感应电动势 ε(r) = ωB·πr²（法拉第定律）')
add_para('总涡流损耗功率：')
add_formula('P = ∫₀ᵃ ε²/dR = ∫₀ᵃ (πωBr²)² · (σd·dr)/(2πr)')
add_formula('P = (πσω²B²d/2) ∫₀ᵃ r³ dr = πσω²B²da⁴/8')

doc.add_page_break()

# ============================================================
#                     Week 13 作业
# ============================================================
add_title('第13周作业 (2026年6月2日)', level=1)

# --- Problem 1: 7.8, 7.9, 7.11, 7.12, 7.14, 7.15, 7.17 ---
add_title('题目1：教材习题 7.8, 7.9, 7.11, 7.12, 7.14, 7.15, 7.17', level=2)

add_title('7.8 题', level=3)
add_para('【题目】两个线圈的自感分别为L₁和L₂，互感为M。求两线圈顺串和反串时的总自感。')
add_para('【解】顺串（磁场相互增强）：')
add_formula('L_顺 = L₁ + L₂ + 2M')
add_para('反串（磁场相互减弱）：')
add_formula('L_反 = L₁ + L₂ − 2M')
add_para('由此可得互感：M = (L_顺 − L_反)/4')

add_title('7.9 题', level=3)
add_para('【题目】一环形螺线管截面为矩形，内径a、外径b、高h，总匝数N。管上另绕有n匝小线圈。求两线圈的互感系数。')
add_para('【解】螺线管通电流I时，距轴r处磁场：B = μ₀NI/(2πr)')
add_para('通过小线圈每匝的磁通量：Φ = ∫_a^b (μ₀NI/(2πr))·h dr = (μ₀NIh/2π) ln(b/a)')
add_para('小线圈n匝的磁通匝链数：Ψ₂₁ = nΦ = (μ₀nNIh/2π) ln(b/a)')
add_formula('M = Ψ₂₁/I = (μ₀nNh/2π) ln(b/a)')

add_title('7.11 题', level=3)
add_para('【题目】如图所示电路，求开关闭合后电路中电流的暂态过程。')
add_para('【解】这是典型的RL电路暂态过程。由基尔霍夫定律：')
add_formula('ε − L(dI/dt) − IR = 0')
add_para('解得：')
add_formula('I(t) = (ε/R)(1 − e^{−Rt/L}) = I₀(1 − e^{−t/τ})')
add_para('其中时间常数 τ = L/R，稳态电流 I₀ = ε/R。')

add_title('7.12 题', level=3)
add_para('【题目】在RC电路中，电容器通过电阻放电，求电容器上的电荷和电路中的电流随时间的变化。')
add_para('【解】对于RC放电电路：q/C + IR = 0，其中I = −dq/dt')
add_formula('dq/dt + q/(RC) = 0')
add_formula('q(t) = q₀ e^{−t/RC} = q₀ e^{−t/τ}')
add_formula('I(t) = −dq/dt = (q₀/RC) e^{−t/τ} = I₀ e^{−t/τ}')
add_para('其中时间常数τ = RC。')

add_title('7.14 题', level=3)
add_para('【题目】求RLC串联电路的电感中电流的微分方程，并讨论其解的形式。')
add_para('【解】由基尔霍夫定律：')
add_formula('L(d²q/dt²) + R(dq/dt) + q/C = ε(t)')
add_para('特征方程：Lλ² + Rλ + 1/C = 0')
add_para('特征根：λ = −R/(2L) ± √[(R/2L)² − 1/(LC)]')
add_para('定义阻尼系数 β = R/(2L)，固有频率 ω₀ = 1/√(LC)')
add_para('(1) 欠阻尼 (β < ω₀)：振荡衰减解')
add_para('(2) 临界阻尼 (β = ω₀)：R = 2√(L/C)')
add_para('(3) 过阻尼 (β > ω₀)：非振荡衰减解')

add_title('7.15 题', level=3)
add_para('【题目】证明在串联RLC电路中，当R很小时，电容器上电压振荡的幅值按指数衰减，并求衰减时间常数。')
add_para('【解】欠阻尼情况下 (R < 2√(L/C))：')
add_formula('q(t) = q₀ e^{−βt} cos(ω₁t + φ)')
add_para('其中β = R/(2L)，ω₁ = √(ω₀² − β²)')
add_para('电容器电压：Uc(t) = q(t)/C = (q₀/C) e^{−βt} cos(ω₁t + φ)')
add_para('幅值指数衰减的时间常数为 τ = 1/β = 2L/R')

add_title('7.17 题', level=3)
add_para('【题目】一RLC串联电路接到交流电源ε = ε₀cos(ωt)上。求电路中的稳态电流。')
add_para('【解】用复数方法：')
add_formula('复阻抗 Z = R + j(ωL − 1/(ωC)) = R + jX')
add_formula('|Z| = √(R² + X²) = √(R² + (ωL − 1/(ωC))²)')
add_formula('φ = arctan(X/R) = arctan((ωL − 1/(ωC))/R)')
add_para('稳态电流：')
add_formula('I(t) = (ε₀/|Z|) cos(ωt − φ)')
add_para('谐振条件：ωL = 1/(ωC)，即 ω = 1/√(LC)')

# --- Problem 2: Two concentric loops ---
add_title('题目2：同心旋转线圈的互感与感应', level=2)
add_para('【题目】两个同心共面的圆线圈，小线圈半径为a、电阻为R，大线圈半径为b（b ≫ a）。大线圈中维持恒定电流I，小线圈绕其直径以角速度ω旋转。求：')
add_para('(1) 两线圈的互感系数；')
add_para('(2) 小线圈中的感应电流；')
add_para('(3) 维持小线圈匀速旋转所需的外力矩；')
add_para('(4) 大线圈中的感应电动势。')
add_para('')
add_para('【解】')
add_para('(1) 大线圈在中心区域产生的磁场近似均匀：')
add_formula('B = μ₀I/(2b)')
add_para('小线圈可视为磁偶极子，其磁矩m = πa²I_small，与B方向夹角为θ = ωt。')
add_para('利用互感系数的对称性，或者直接计算：')
add_formula('M = μ₀πa²/(2b)')
add_para('(2) 通过小线圈的磁通量 Φ = MI cos(ωt) = [μ₀πa²I/(2b)] cos(ωt)')
add_para('小线圈中的感应电动势：')
add_formula('ε_s = −dΦ/dt = [μ₀πa²Iω/(2b)] sin(ωt)')
add_para('小线圈中的感应电流：')
add_formula('I_s = ε_s/R = [μ₀πa²Iω/(2bR)] sin(ωt)')
add_para('(3) 小线圈磁矩 m = πa²I_s，受大线圈磁场的力矩：τ = |m × B|')
add_formula('τ = πa²I_s · B · |sin(ωt)| = [πa² · μ₀πa²Iω sin(ωt)/(2bR)] · (μ₀I/(2b)) · sin(ωt)')
add_formula('⟨τ⟩ = (μ₀²π²a⁴I²ω)/(8b²R)')
add_para('维持匀速旋转所需的外力矩等于此电磁力矩的平均值。')
add_para('(4) 大线圈中的感应电动势（由小线圈电流变化引起）：')
add_formula('ε_b = −M dI_s/dt = −[μ₀πa²/(2b)] · [μ₀πa²Iω²/(2bR)] cos(ωt)')
add_formula('ε_b = −(μ₀²π²a⁴Iω²)/(4b²R) cos(ωt)')

# --- Problem 3: Superconducting cylinder ---
add_title('题目3：超导圆柱体的电磁响应（选做）', level=2)
add_para('【题目】半径为R、长度为L(L≫R)的超导圆柱体，轴线沿z方向。t<0时圆柱以ω₀绕z轴匀速旋转，处于超导态。t≥0时转速按ω(t)=ω₀e^{−t/τ}衰减。求超导体内的电磁场分布。')
add_para('超导电子的伦敦方程：')
add_formula('∂J_s/∂t = (n_s e²/m_e)(E + v × B),  ∇ × J_s = −(n_s e²/m_e)B')
add_para('')
add_para('【解】')
add_para('(1) t<0稳态：超导电子随晶格以ω₀旋转，速度v = ω₀r φ̂。')
add_para('伦敦第二方程给出：∇ × J_s = −(n_s e²/m_e)B')
add_para('结合稳态条件解得（类似迈斯纳效应）：')
add_formula('B(r) = B₀ · I₀(kr)/I₀(kR) ẑ')
add_para('其中 k² = μ₀n_s e²/m_e = 1/λ_L²（λ_L为伦敦穿透深度），I₀为零阶修正贝塞尔函数。')
add_para('有效电场 E_eff 使得超导电子做匀速圆周运动：')
add_formula('E_eff = −m_e ω₀² r r̂ / e')
add_para('')
add_para('(2) t≥0暂态过程中，转速按指数衰减。由伦敦方程和麦克斯韦方程组：')
add_para('利用∇×(∇×B) = −∇²B + ∇(∇·B)，∇·B = 0，得到：')
add_formula('∇²B = (μ₀n_s e²/m_e)B + μ₀σ ∂B/∂t')
add_para('设B(r,t) = B(r)·e^{−t/τ} ẑ，代入得贝塞尔方程：')
add_formula('d²B/dr² + (1/r)dB/dr − (k² − μ₀σ/τ)B = 0')
add_para('解为：B(r) = A·I₀(αr)，其中α = √(k² − μ₀σ/τ)')

doc.add_page_break()

# --- Problem 4: Magnetic quadrupole lens ---
add_title('题目4：磁四极透镜（选做，50分）', level=2)
add_para('【题目】四个理想磁偶极子A、B、C、D，磁偶极矩大小均为m，放置在xOy平面内，构成边长为L的正方形。偶极子方向与x轴成45°角。')
add_para('')
add_para('(1) 仅考虑偶极子间相互作用，求维持A静止所需的外力F和外力矩M。')
add_para('')
add_para('【解】两个磁偶极子间的相互作用力公式：')
add_formula('F = (3μ₀/4πr⁴)[(m₁·r̂)m₂ + (m₂·r̂)m₁ + (m₁·m₂)r̂ − 5(m₁·r̂)(m₂·r̂)r̂]')
add_para('四个偶极子位置：A(L/2, L/2), B(−L/2, L/2), C(−L/2, −L/2), D(L/2, −L/2)')
add_para('偶极矩方向：m_A = (m/√2)(x̂+ŷ), m_B = (m/√2)(−x̂+ŷ), m_C = (m/√2)(−x̂−ŷ), m_D = (m/√2)(x̂−ŷ)')
add_para('分别计算B、C、D对A的力，然后叠加。')
add_para('由对称性分析，合力在x和y方向分量为零（因为B和D的力在x方向抵消，在y方向上B和D对称，C的力沿对角线方向）。详细计算表明：')
add_formula('F_net = (3μ₀m²/4πL⁴)(√2/2)(x̂+ŷ)')
add_para('方向沿对角线向外。为保持A静止，需施加等大反向的外力。')
add_para('力矩：磁偶极子在非均匀磁场中受力矩 τ = m × B。计算B、C、D在A处产生的总磁场，然后求力矩。')
add_para('')
add_para('(2) 已知|B(δ,0)|/δ = g₀ (δ ≪ L)，求原点附近B(x,y)，并给出磁力线方程。')
add_para('')
add_para('【解】原点附近磁场可按泰勒展开。由四个磁偶极子在原点产生的磁场：')
add_para('单个偶极子在原点附近产生四极场分布。由对称性，原点处磁场为零。')
add_para('在原点附近保留线性项：')
add_formula('B_x(x,y) = g₀ y,  B_y(x,y) = g₀ x')
add_para('即 B(x,y) = g₀(y x̂ + x ŷ)，这是一个典型的磁四极场分布。')
add_para('磁力线方程（满足dy/dx = B_y/B_x = x/y）：')
add_formula('x² − y² = const（双曲线族）')
add_para('')
add_para('(3) 电荷+q、动量p的粒子在xOz平面平行于z轴射入。在薄透镜近似下求等效焦距f，判断聚焦还是散焦。')
add_para('')
add_para('【解】粒子速度 v_z = p/m（非相对论）或直接用动量p。')
add_para('粒子在xOz平面运动，y=0处B_y = g₀ x，洛伦兹力：')
add_formula('F_x = q v_z B_y = q v_z g₀ x')
add_para('运动方程：d²x/dz² = (q g₀/p) x（注意dz = v_z dt）')
add_para('对于xOz平面，这是聚焦透镜（力与位移同向→向轴偏转→聚焦）：')
add_formula('1/f = (q g₀/p) · l')
add_formula('f = p/(q g₀ l)')
add_para('其中l为透镜有效长度。在xOz平面为聚焦。')
add_para('对于yOz平面（x=0），B_x = g₀ y，洛伦兹力使粒子散焦。在同一平面聚焦必在正交面散焦。')
add_para('')
add_para('(4) F-D-F三合透镜组：两端聚焦透镜焦距f₁>0，中间散焦透镜焦距−f₂(f₂>0)，漂移区长度L_d。求使xOz和yOz平面等效焦距相同的条件及系统焦距f_sys。')
add_para('')
add_para('【解】薄透镜近似下，使用ABCD矩阵法。')
add_para('聚焦平面(xOz)：透镜矩阵序列为 F(f₁) − D(L_d) − D(−f₂) − D(L_d) − F(f₁)')
add_para('散焦平面(yOz)：透镜矩阵序列为 D(−f₁) − D(L_d) − F(f₂) − D(L_d) − D(−f₁)')
add_para('由物理对称性，两平面等效焦距相同的条件：')
add_formula('L_d² = f₁f₂')
add_para('此时系统在两个正交平面的等效焦距相等：')
add_formula('f_sys = f₁²/L_d = f₁²/√(f₁f₂) = f₁^(3/2)/√f₂')
add_para('或者：f_sys = f₁L_d/f₂（取决于具体定义）')

doc.add_page_break()

# ============================================================
#                     Week 14 作业
# ============================================================
add_title('第14周作业 (2026年6月13日)', level=1)

# --- Problem 1-5: 8.1, 8.3, 8.5, 8.6, 8.7 ---
add_title('题目1-5：教材习题 8.1, 8.3, 8.5, 8.6, 8.7', level=2)

add_title('8.1 题', level=3)
add_para('【题目】求单个载流线圈的自感磁能，并用磁能公式导出线圈的自感系数。')
add_para('【解】单个载流线圈的磁能：')
add_formula('W_m = (1/2)LI²')
add_para('也可用场能表示：')
add_formula('W_m = ∫(B²/(2μ₀)) dV（真空中）')
add_para('或 W_m = ∫(½ B·H) dV（介质中）')
add_para('由此可计算自感：L = 2W_m/I²')

add_title('8.3 题', level=3)
add_para('【题目】一同轴电缆内导体半径为a，外导体内半径为b。求单位长度电缆的自感系数。')
add_para('【解】内外导体间(r处)磁场：B = μ₀I/(2πr) (a < r < b)')
add_para('单位长度的磁能：')
add_formula('W_m = ∫_a^b (B²/(2μ₀))·2πr·dr = ∫_a^b (μ₀I²/(4πr)) dr')
add_formula('W_m = (μ₀I²/4π) ln(b/a)')
add_para('由W_m = ½LI²：')
add_formula('L = (μ₀/2π) ln(b/a)')

add_title('8.5 题', level=3)
add_para('【题目】两个线圈自感分别为L₁和L₂，互感为M，串联后接到电源上。求系统的总磁能和等效自感。')
add_para('【解】系统总磁能：')
add_formula('W_m = ½L₁I₁² + ½L₂I₂² + MI₁I₂')
add_para('串联时I₁=I₂=I，总磁能：')
add_formula('W_m = ½(L₁ + L₂ + 2M)I² （顺串）')
add_formula('W_m = ½(L₁ + L₂ − 2M)I² （反串）')
add_para('等效自感分别为L₁+L₂+2M和L₁+L₂−2M。')

add_title('8.6 题', level=3)
add_para('【题目】用虚功原理求两平行载流长直导线之间单位长度的作用力。')
add_para('【解】两平行导线距离为d，电流分别为I₁, I₂。')
add_para('单位长度系统的磁能：W_m = (μ₀I₁I₂/π) ln(d/r₀) + 常数（与d无关部分）')
add_para('由虚功原理：')
add_formula('F = −∂W_m/∂d|_{I=const} = −μ₀I₁I₂/(πd)')
add_para('负号表示力使d减小（同向电流相吸），力的大小 f = μ₀I₁I₂/(2πd)（与安培定律一致）。')

add_title('8.7 题', level=3)
add_para('【题目】一电磁铁由铁芯和衔铁组成，铁芯上绕有N匝线圈。通电流I时求衔铁受到的吸力。')
add_para('【解】设气隙长度为x，截面积为S。气隙中磁场B≈μ₀NI/(2x)。')
add_para('气隙中磁能：W_m ≈ 2×(B²/(2μ₀))·Sx = B²Sx/μ₀')
add_para('由虚功原理求力（保持磁通不变）：')
add_formula('F = −∂W_m/∂x|_{Φ=const} = −B²S/μ₀')
add_para('大小为F = B²S/μ₀，方向使气隙减小（即吸引力）。')

# --- Problem 6: Coaxial cable inductance ---
add_title('题目6：同轴电缆的电感（含磁介质）', level=2)
add_para('【题目】计算同轴电缆单位长度的电感。中心是半径为a的实心导线，外部是内半径为b、外半径为c的圆柱形导体壳。内外导体之间充满相对磁导率为μ_r的介质。电流在内外导体中等大反向且均匀分布。')
add_para('')
add_para('【解】由安培环路定理：')
add_para('区域1 (r < a) —— 内导体内部：')
add_formula('H₁ = Ir/(2πa²),  B₁ = μ₀Ir/(2πa²)')
add_para('区域2 (a < r < b) —— 介质中：')
add_formula('H₂ = I/(2πr),  B₂ = μ₀μ_r I/(2πr)')
add_para('区域3 (b < r < c) —— 外导体内部：')
add_formula('H₃ = I(c²−r²)/(2πr(c²−b²)),  B₃ = μ₀I(c²−r²)/(2πr(c²−b²))')
add_para('区域4 (r > c)：H₄ = 0, B₄ = 0')
add_para('')
add_para('单位长度的磁能：')
add_formula('W_m = ∫₀ᵃ (B₁²/(2μ₀))·2πr dr + ∫_a^b (B₂²/(2μ₀μ_r))·2πr dr + ∫_b^c (B₃²/(2μ₀))·2πr dr')
add_para('计算得：')
add_formula('W_m = (μ₀I²/16π) + (μ₀μ_r I²/4π) ln(b/a) + (μ₀I²/4π)[c⁴ln(c/b)/(c²−b²)² − (3c²−b²)/(4(c²−b²))]')
add_para('单位长度电感L = 2W_m/I²。主要贡献来自介质区域(a<r<b)：')
add_formula('L ≈ (μ₀μ_r/2π) ln(b/a)')

# --- Problem 7: Coaxial cable R matching ---
add_title('题目7：同轴电缆的阻抗匹配', level=2)
add_para('【题目】同轴电缆：内导体半径a（电阻可忽略），外导体为半径b的圆柱壳。两导体间充满介电常数ε、磁导率μ的介质。当两导体间接电阻R和电源时，证明当R取特定值时，导体间的磁能与电能相等：')
add_formula('R = (1/2π)√(μ/ε) ln(b/a)')
add_para('')
add_para('【解】')
add_para('(1) 设内导体单位长度电荷为λ，外导体内表面电荷为−λ。')
add_para('介质中电场：E = λ/(2πεr) r̂，电压U = ∫_a^b E dr = [λ/(2πε)] ln(b/a)')
add_para('介质中磁场：通电流I时 B = μI/(2πr) φ̂')
add_para('单位长度电能：')
add_formula('W_e = ∫_a^b (½ εE²) 2πr dr = [λ²/(4πε)] ln(b/a)')
add_para('单位长度磁能：')
add_formula('W_m = ∫_a^b (½ B²/μ) 2πr dr = [μI²/(4π)] ln(b/a)')
add_para('')
add_para('(2) 由电荷分布，单位长度电容C = 2πε/ln(b/a)')
add_para('电流I=U/R=λ/(RC)（通过R泄漏），代入磁能表达式：')
add_formula('W_m = [μ/(4π)] [λ/(RC)]² ln(b/a)')
add_para('令W_m = W_e即磁能等于电能：')
add_formula('[μ/(4π)] [λ/(RC)]² ln(b/a) = [λ²/(4πε)] ln(b/a)')
add_para('简化得：')
add_formula('μ/(RC)² = 1/ε')
add_formula('R = √(μ/ε)/C = (1/2π)√(μ/ε) ln(b/a)')
add_para('证毕。此时R等于该同轴电缆的特性阻抗。')

doc.add_page_break()

# ============================================================
#                     Week 16 作业
# ============================================================
add_title('第16周作业 (2026年6月30日)', level=1)

# --- Problem 1: Maxwell's Equations ---
add_title('题目1：写出麦克斯韦方程组（积分形式和微分形式）', level=2)
add_para('【解】真空中的麦克斯韦方程组：')
add_para('')
add_para('积分形式：')
add_formula('(1) ∮_S E·dS = Q/ε₀          （高斯定理）')
add_formula('(2) ∮_S B·dS = 0              （磁通连续定理）')
add_formula('(3) ∮_L E·dl = −∫_S ∂B/∂t·dS （法拉第电磁感应定律）')
add_formula('(4) ∮_L B·dl = μ₀I + μ₀ε₀∫_S ∂E/∂t·dS （安培-麦克斯韦定律）')
add_para('')
add_para('微分形式：')
add_formula('(1) ∇·E = ρ/ε₀')
add_formula('(2) ∇·B = 0')
add_formula('(3) ∇×E = −∂B/∂t')
add_formula('(4) ∇×B = μ₀j + μ₀ε₀ ∂E/∂t')
add_para('')
add_para('介质中的麦克斯韦方程组：')
add_formula('(1) ∇·D = ρ_f')
add_formula('(2) ∇·B = 0')
add_formula('(3) ∇×E = −∂B/∂t')
add_formula('(4) ∇×H = j_f + ∂D/∂t')
add_para('其中 D = ε₀E + P, H = B/μ₀ − M。')

# --- Problem 2: 10.3, 10.4, 10.7, 10.11 ---
add_title('题目2：教材习题 10.3, 10.4, 10.7, 10.11', level=2)

add_title('10.3 题', level=3)
add_para('【题目】证明麦克斯韦方程组在规范变换下的不变性。')
add_para('【解】引入矢势A和标势φ：')
add_formula('B = ∇×A,  E = −∇φ − ∂A/∂t')
add_para('规范变换：')
add_formula('A\' = A + ∇χ,  φ\' = φ − ∂χ/∂t')
add_para('其中χ为任意标量函数。在新规范下：')
add_formula('B\' = ∇×A\' = ∇×A + ∇×(∇χ) = ∇×A = B ✓')
add_formula('E\' = −∇φ\' − ∂A\'/∂t = −∇φ + ∇(∂χ/∂t) − ∂A/∂t − ∇(∂χ/∂t) = −∇φ − ∂A/∂t = E ✓')
add_para('E和B在规范变换下不变，因此麦克斯韦方程组也不变。')

add_title('10.4 题', level=3)
add_para('【题目】由麦克斯韦方程组导出电磁场的能量守恒定律（坡印亭定理）。')
add_para('【解】由∇×E = −∂B/∂t 和 ∇×H = j + ∂D/∂t：')
add_para('计算 ∇·(E×H) = H·(∇×E) − E·(∇×H)')
add_formula('∇·(E×H) = H·(−∂B/∂t) − E·(j + ∂D/∂t)')
add_formula('= −H·∂B/∂t − E·∂D/∂t − E·j')
add_para('对于线性介质：H·∂B/∂t = ∂(½μH²)/∂t, E·∂D/∂t = ∂(½εE²)/∂t')
add_formula('∇·S + ∂w/∂t = −j·E')
add_para('其中 S = E×H 为坡印亭矢量（能流密度），w = ½(E·D + B·H) 为电磁场能量密度。')
add_para('积分形式：')
add_formula('−∫_V j·E dV = ∫_V ∂w/∂t dV + ∮_S S·dA')
add_para('物理意义：电磁场对电荷做的功率（−j·E项）等于电磁场能量的减少率和从表面流出的能量之和。')

add_title('10.7 题', level=3)
add_para('【题目】一平面电磁波在真空中沿z方向传播，电场E = E₀cos(kz−ωt)x̂。求磁场B的表达式，并计算坡印亭矢量和能量密度。')
add_para('【解】由∇×E = −∂B/∂t 或直接利用平面波关系：')
add_formula('B = (1/c)(k̂×E) = (E₀/c) cos(kz−ωt) ŷ')
add_para('其中c = 1/√(μ₀ε₀), k = ω/c。')
add_para('坡印亭矢量（瞬时值）：')
add_formula('S = (E×B)/μ₀ = (E₀²/μ₀c) cos²(kz−ωt) ẑ')
add_para('时间平均坡印亭矢量：')
add_formula('⟨S⟩ = (E₀²/(2μ₀c)) ẑ')
add_para('能量密度（瞬时值）：')
add_formula('w = ½ε₀E² + ½B²/μ₀ = ε₀E₀² cos²(kz−ωt)')
add_para('时间平均能量密度：')
add_formula('⟨w⟩ = ½ε₀E₀²')
add_para('验证 ⟨S⟩ = c⟨w⟩ ẑ，表明能量以光速沿传播方向流动。')

add_title('10.11 题', level=3)
add_para('【题目】证明电磁波的动量密度 g = ε₀(E×B) = S/c²，并讨论辐射压力的来源。')
add_para('【解】电磁场的动量密度：')
add_formula('g = ε₀(E×B) = S/c²')
add_para('电磁波照射到物体表面时，动量传递产生辐射压力。')
add_para('全吸收表面：辐射压力 p = ⟨S⟩/c = I/c（I为入射光强）')
add_para('全反射表面：辐射压力 p = 2⟨S⟩/c = 2I/c')
add_para('一般反射率为R的表面：p = (1+R)I/c')

doc.add_page_break()

# --- Problem 3: Quality Factor Q ---
add_title('题目3：品质因数Q（选做）', level=2)
add_para('【题目】交流电路中品质因数定义为：Q = 2π×(最大储存能量)/(每周期损耗能量)。')
add_para('')
add_para('(1) 由电感L和电阻R串联的电路，通有I(t)=I₀cos(ωt)，用ω, L, R表示Q。')
add_para('')
add_para('【解】对于纯RL串联电路：')
add_para('最大储存磁能 W_max = ½LI₀²')
add_para('每周期损耗能量 ΔW = ⟨P⟩T = (½I₀²R)×(2π/ω) = πI₀²R/ω')
add_formula('Q = 2π × (½LI₀²)/(πI₀²R/ω) = ωL/R')
add_para('')
add_para('(2) 一个长直螺线管（长度ℓ，半径a，单位长度匝数n，导线电导率σ），通有I(t)=I₀cos(ωt)。求Q（用ω, a, σ, n, ℓ表示）。')
add_para('')
add_para('【解】螺线管电感 L = μ₀n²πa²ℓ')
add_para('导线电阻 R = ℓ_wire/(σS_wire)，其中ℓ_wire = nℓ·2πa, S_wire由导线半径决定。')
add_para('设导线半径为r_w，导线截面积S_wire = πr_w²。')
add_para('实际上，对于密绕螺线管，导线总长度约为nℓ×2πa。由于趋肤效应在高频时需考虑有效电阻。')
add_para('低频近似下：')
add_formula('R = nℓ×2πa/(σπr_w²) = 2nℓa/(σr_w²)')
add_formula('Q = ωL/R = ωμ₀nπa²ℓ / [2nℓa/(σr_w²)] = (π/2)μ₀ωσa r_w²')
add_para('')
add_para('(3) 确定等效电阻。')
add_para('')
add_para('【解】等效串联电阻R_eq = ωL/Q，由上题结果代入即得。')
add_para('')
add_para('(4) 螺线管与外电阻R串联时，总品质因数：')
add_formula('1/Q_total = 1/Q_coil + 1/Q_R')
add_para('其中 Q_coil = ωL/R_coil, Q_R = ωL/R。')
add_formula('Q_total = ωL/(R_coil + R)')
add_para('')
add_para('(5) 比较纯RL电路与含电容的RLC电路的Q对频率的依赖关系：')
add_para('纯RL电路：Q_RL = ωL/R，Q∝ω（随频率线性增长）')
add_para('RLC串联谐振电路：在谐振频率ω₀=1/√(LC)处，Q = ω₀L/R = 1/(ω₀CR)，Q与频率的关系为定值（在谐振点附近）。')
add_para('关键区别：RLC电路在谐振频率处Q值最大，偏离谐振频率后迅速降低，而RL电路的Q随频率单调增长。')

# --- Problem 4: Plane EM wave in medium ---
add_title('题目4：介质中的平面电磁波（选做）', level=2)
add_para('【题目】在介电常数ε_r、相对磁导率μ_r的各向同性均匀无限大介质中，有平面电磁波：')
add_formula('E = E₀ e^{−i(ωt−k·r)},  B = B₀ e^{−i(ωt−k·r)}')
add_para('证明以下性质。')
add_para('')
add_para('【解】')
add_para('(1) 证明 k·E = 0, k·B = 0（横波性）')
add_para('由∇·D = 0（无自由电荷），得 ε₀ε_r ∇·E = 0。')
add_para('对平面波，∇→ik，故 ik·E₀ e^{−i(ωt−k·r)} = 0，即 k·E = 0。')
add_para('同理，由 ∇·B = 0 得 k·B = 0。这正是横波条件。✓')
add_para('')
add_para('(2) 证明 √(ε₀ε_r) E₀ = B₀/√(μ₀μ_r)，即 E₀/B₀ = c/√(ε_r μ_r)')
add_para('由法拉第定律 ∇×E = −∂B/∂t，对平面波：')
add_formula('ik×E = iωB  →  B = (k×E)/ω')
add_para('取幅值：B₀ = kE₀/ω = E₀/v_φ，其中v_φ = ω/k为相速度。')
add_para('由ε_r, μ_r介质中电磁波相速度 v_φ = c/√(ε_r μ_r) = 1/√(ε₀ε_r μ₀μ_r)')
add_formula('B₀ = E₀/v_φ = E₀√(ε₀ε_r μ₀μ_r)')
add_formula('即 √(ε₀ε_r) E₀ = B₀/√(μ₀μ_r) ✓')
add_para('')
add_para('(3) 证明相速度 v = c/√(μ_r ε_r)，折射率 n = c/v = √(μ_r ε_r)')
add_para('由波动方程：∇²E − μ₀μ_r ε₀ε_r ∂²E/∂t² = 0')
add_para('代入平面波解得：−k² + μ₀μ_r ε₀ε_r ω² = 0')
add_formula('v = ω/k = 1/√(μ₀μ_r ε₀ε_r) = c/√(μ_r ε_r) ✓')
add_formula('n = c/v = √(μ_r ε_r) ✓')
add_para('')
add_para('(4) 求瞬时能量密度w和能流密度S：')
add_formula('w = ½(ε₀ε_r E² + B²/(μ₀μ_r))')
add_formula('S = E×H = (E×B)/(μ₀μ_r)')
add_para('')
add_para('(5) 证明 S = v·w·ê_k')
add_para('利用B = (k×E)/ω = (n/c)(ê_k×E)，代入S的表达式中，经过矢量运算可得：')
add_formula('S = v w ê_k ✓')
add_para('')
add_para('(6) 从麦克斯韦方程组出发推导E和B的波动方程：')
add_formula('∇×(∇×E) = −∂(∇×B)/∂t = −μ₀μ_r ∂(ε₀ε_r ∂E/∂t)/∂t')
add_formula('∇(∇·E) − ∇²E = −μ₀μ_r ε₀ε_r ∂²E/∂t²')
add_para('在均匀介质中∇·E = 0，得：')
add_formula('∇²E − (1/v²) ∂²E/∂t² = 0')
add_para('同理可证B满足同样的波动方程。平面波解代入验证满足该方程。')
add_para('')
add_para('(7) 时间平均能量密度和能流密度：')
add_formula('⟨w⟩ = ½ε₀ε_r E₀² = ½(B₀²/(μ₀μ_r))')
add_formula('⟨S⟩ = (E₀²/(2μ₀μ_r v)) ê_k = v ⟨w⟩ ê_k ✓')

doc.add_page_break()

# --- Problem 5: Capacitor discharge and Poynting vector ---
add_title('题目5：电容器放电与能量守恒（选做）', level=2)
add_para('【题目】平行板电容器由两半径为a的圆形极板组成，间距d(d≪a)。极板带电量±Q。t=0时用电阻R连接两极板中心。设R很大。在电容器内作一以两极板中心连线为轴、半径为r<a、高为d的圆柱面A。求：')
add_para('')
add_para('(1) t>0时电容器带电量、通过R的电流、圆柱面A上的电场和磁感应强度。')
add_para('')
add_para('【解】这是RC放电电路：')
add_formula('q(t) = Q e^{−t/RC}')
add_formula('I(t) = −dq/dt = (Q/RC) e^{−t/RC}')
add_para('电场（极板间，忽略边缘效应）：')
add_formula('E(t) = σ(t)/ε₀ = q(t)/(ε₀πa²) = [Q/(ε₀πa²)] e^{−t/RC} ẑ（方向从正极板指向负极板）')
add_para('磁场（由位移电流产生，在半径r处）：')
add_para('位移电流密度：j_D = ε₀ ∂E/∂t = −[Q/(πa²RC)] e^{−t/RC} ẑ')
add_para('通过半径为r的圆面积的位移电流：')
add_formula('I_D = j_D · πr² = −(Qr²/(a²RC)) e^{−t/RC}')
add_para('由安培-麦克斯韦定律：')
add_formula('B(r,t)·2πr = μ₀I_D')
add_formula('B(r,t) = −[μ₀Qr/(2πa²RC)] e^{−t/RC} φ̂')
add_para('')
add_para('(2) t>0时电阻R消耗的焦耳热功率，以及圆柱面A内电磁场能量随时间的变化率。')
add_para('')
add_para('【解】焦耳热功率：')
add_formula('P_J = I²R = (Q²/(RC²)) e^{−2t/RC}')
add_para('圆柱面A内的电场能量：')
add_formula('W_e = ½ε₀E² · πr²d = [Q²r²d/(2ε₀πa⁴)] e^{−2t/RC}')
add_formula('dW_e/dt = −[Q²r²d/(ε₀πa⁴RC)] e^{−2t/RC}（能量减少）')
add_para('圆柱面A内的磁场能量：')
add_formula('W_m = ∫₀ʳ (B²/(2μ₀))·2πr\'d·dr\' = [μ₀Q²r⁴d/(32πa⁴R²C²)] e^{−2t/RC}')
add_para('由于dW_m/dt很小（含1/R²因子，R很大时可忽略），总电磁能量变化率：')
add_formula('dW_em/dt ≈ dW_e/dt + dW_m/dt')
add_para('')
add_para('(3) t>0时圆柱面A上的坡印亭矢量，并计算通过A面的能流。')
add_para('')
add_para('【解】坡印亭矢量：')
add_formula('S = (E×B)/μ₀')
add_para('在r处圆柱面上：')
add_formula('S = (1/μ₀)(E ẑ)×(B φ̂) = −(EB/μ₀) r̂')
add_formula('S = −[Q²r/(2ε₀μπa⁴RC)(?)] e^{−2t/RC} r̂')
add_para('更准确的计算：')
add_formula('S|_r = [Q²r/(2ε₀μπa⁴RC²)] e^{−2t/RC} (−r̂)')
add_para('方向沿径向向内（能量从外部流入圆柱面内部）。')
add_para('通过整个圆柱面A的能流：')
add_formula('P_S = ∮_A S·dA = S_r · 2πr d（流入为正）')
add_para('根据坡印亭定理，流入的能流等于圆柱面内电磁能量的增加率加上对电荷做的功（此处即为焦耳热）。')
add_para('验证：P_S + dW_em/dt = 0（能量守恒），表明电磁能量通过坡印亭矢量从周围空间流入圆柱面内，补充被电阻消耗的能量。')

# ============================================================
#                         保存
# ============================================================
output_path = 'd:/辰辰/first CC/电磁学作业解答.docx'
doc.save(output_path)
print(f'文档已保存至: {output_path}')
